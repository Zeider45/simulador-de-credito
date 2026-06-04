#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera la documentacion Word del funcionamiento de los creditos UVCC."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/user/simulador-de-credito/docs/Funcionamiento-Creditos-UVCC.docx"

# ---- Colores ----
AZUL = RGBColor(0x1F, 0x3A, 0x5F)
GRIS = RGBColor(0x60, 0x60, 0x60)
VERDE = RGBColor(0x1B, 0x5E, 0x20)
NARANJA = RGBColor(0x9A, 0x52, 0x00)
MORADO = RGBColor(0x4A, 0x14, 0x8C)

doc = Document()

# Estilo base legible
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = AZUL
    return h

def source(kind, text):
    """kind: 'doc' (documental), 'code' (codigo), 'own' (elaboracion propia)."""
    cfg = {
        'doc':  ("Fuente documental", VERDE),
        'code': ("Base en el codigo del proyecto", AZUL),
        'own':  ("Elaboracion propia (no consta en la documentacion; creado por el asistente)", NARANJA),
    }
    label, color = cfg[kind]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"▶ {label}: ")
    r.bold = True; r.italic = True; r.font.size = Pt(9); r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = GRIS
    return p

def para(text, bullet=False):
    p = doc.add_paragraph(text, style="List Bullet" if bullet else None)
    return p

def formula(lines):
    """Bloque monoespaciado simple para formulas."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"; run.font.size = Pt(10); run.font.color.rgb = MORADO
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shade(hdr[i], "1F3A5F")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============================ PORTADA ============================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Funcionamiento de los Creditos Comerciales\nexpresados en Unidad de Valor de Credito (UVC/UVCC)")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Formulas, asientos contables, clasificacion de cartera y situaciones operativas\nDocumentacion tecnica y regulatoria del simulador de credito")
rs.font.size = Pt(12); rs.font.color.rgb = GRIS
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
rd = d.add_run("Version 1.0"); rd.font.size = Pt(10); rd.font.color.rgb = GRIS

doc.add_paragraph()

# ---- Leyenda de fuentes ----
heading("Como leer este documento", 2)
para("Bajo el titulo de cada seccion se indica el origen de la informacion mediante una de estas tres etiquetas:")
lt = doc.add_table(rows=0, cols=2); lt.style = "Light List Accent 1"
for lab, desc, col in [
    ("▶ Fuente documental", "La informacion proviene de un archivo de la carpeta 'documentacion/' (resoluciones BCV, circulares SUDEBAN, presentacion contable o modelos Excel). Se cita el documento.", VERDE),
    ("▶ Base en el codigo del proyecto", "La informacion describe como esta implementado el calculo en el codigo (src/lib/simulator.js y regulatory.js).", AZUL),
    ("▶ Elaboracion propia", "La informacion NO consta en la documentacion entregada: es un criterio o supuesto creado por el asistente para completar el modelo.", NARANJA),
]:
    cells = lt.add_row().cells
    cells[0].text=""; rr=cells[0].paragraphs[0].add_run(lab); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=col
    cells[1].text=""; r2=cells[1].paragraphs[0].add_run(desc); r2.font.size=Pt(9.5)
    cells[0].width=Inches(2.1); cells[1].width=Inches(4.4)

doc.add_page_break()

# ============================ 1. INTRODUCCION ============================
heading("1. Introduccion y alcance", 1)
source('own', "Sintesis redactada por el asistente a partir del conjunto del proyecto.")
para("Este documento describe como debe funcionar un credito comercial expresado en Unidad de Valor de Credito (UVC), tambien llamada Unidad de Valor de Credito Comercial (UVCC) en la normativa de 2019. Cubre las formulas de calculo, el desglose de la cuota, la mora, la clasificacion de la cartera, los asientos contables, el tratamiento del credito vencido, la cancelacion anticipada y distintas situaciones operativas.")
para("El objetivo es que cualquier persona (negocio, contabilidad o desarrollo) entienda el modelo completo y pueda auditar cada numero contra su fundamento.")

# ============================ 2. MARCO REGULATORIO ============================
heading("2. Marco regulatorio de referencia", 1)
source('doc', "Carpeta 'documentacion/': resoluciones BCV, circulares SUDEBAN y avisos oficiales.")
table(
    ["Instrumento", "Fecha / Gaceta", "Que aporta"],
    [
        ["Resolucion BCV N° 19-09-01", "G.O. 41.742 del 21/10/2019", "Crea la obligacion de expresar los creditos comerciales en UVCC y dividir el monto en Bs entre el Indice de Inversion (IDI)."],
        ["Aviso Oficial BCV", "G.O. 41.742 del 21/10/2019", "Comision flat maxima de 0,50% del monto del credito."],
        ["Circular BCV de entrada en vigencia", "24/10/2019", "Fija la aplicacion de la Resolucion 19-09-01 a partir del 28/10/2019."],
        ["Circular SUDEBAN SIB-DSB-CJ-OD-13083", "14/11/2019", "Clausulas minimas del contrato: amortizacion de capital en UVCC, cancelacion anticipada con piso de IDI, % de mora segun el BCV."],
        ["Modificacion Manual de Contabilidad (Circular SIB-II-GGR-GNP-12161)", "28/10/2019", "Crea las subcuentas contables 131.35, 131.36, 358.01, 513.01.M.35 y 513.01.M.36."],
        ["Minuta reunion SUDEBAN/BCV/ABV", "17/12/2019", "Tratamiento del credito vencido: congelamiento y registro en cuentas de orden/patrimonio."],
        ["Resolucion BCV N° 21-01-02 (VIGENTE)", "G.O. 42.050 del 19/01/2021; vigente 01/02/2021", "Norma actual: UVC, IDI, tasas 4%-10% (2% Cartera Productiva), mora 0,80%, piso de IDI en cancelacion anticipada."],
    ],
    widths=[2.5, 1.7, 2.6],
)

# ============================ 3. UVC e IDI ============================
heading("3. Unidad de Valor de Credito (UVC) e Indice de Inversion (IDI)", 1)
source('doc', "Resolucion BCV 21-01-02, Art. 1; presentacion 'Impacto contable UVCC'.")
para("La UVC es una unidad de cuenta para preservar el valor real del credito. El IDI (Indice de Inversion) es un factor que el BCV publica diariamente y expresa cuantos bolivares vale una UVC en una fecha. El IDI refleja la variacion del tipo de cambio de referencia de mercado.")
para("Conversion del capital al momento del desembolso:")
formula(["Capital_UVC = Capital_Bs / IDI_desembolso"])
para("Conversion de cualquier monto UVC a bolivares en una fecha:")
formula(["Monto_Bs = Monto_UVC x IDI_fecha"])
heading("3.1 Resolucion del IDI en dias sin publicacion", 2)
source('own', "Criterio del sistema: el BCV no fija un metodo obligatorio de relleno. Implementado en el codigo (createIdiResolver).")
para("Fin de semana: se usa el IDI del ultimo dia habil previo.", bullet=True)
para("Fechas pasadas sin dato: arrastre del ultimo IDI conocido o interpolacion lineal (parametro idiMissing).", bullet=True)
para("Fechas futuras: extrapolacion sumando un incremento diario configurable (idiFutureStep).", bullet=True)

# ============================ 4. AMORTIZACION ============================
heading("4. Modalidades de pago y amortizacion", 1)
source('doc', "Resolucion BCV 21-01-02, Art. 5 lit. a y b; modelos Excel (Simulacion UVCC, SIMULACION 30/90/120).")
para("La norma exige que cada cuota incluya interes y una porcion de amortizacion de capital expresada en UVC. Existen dos modalidades documentadas:")
para("Pago en cuotas (sistema frances, cuota fija en UVC): hoja 'ejercicio en cuotas' de 'Simulacion UVCC'.", bullet=True)
para("Pago unico al vencimiento (bullet, 'AL VCTO'): hojas '30/90/120 dias'. Para Cartera Productiva (Art. 2) lleva un cargo especial de 20% al liquidar.", bullet=True)
heading("4.1 Formula de la cuota fija en UVC (sistema frances)", 2)
source('own', "La exigencia de amortizar capital en cada cuota es regulatoria (Art. 5 lit. a); la formula francesa concreta es criterio del sistema.")
formula([
    "i = tasa_anual / 12            (tasa mensual)",
    "Cuota_UVC = Capital_UVC x [ i / (1 - (1 + i)^(-n)) ]",
    "  n = numero de cuotas (meses)",
    "Caso tasa 0:  Cuota_UVC = Capital_UVC / n",
])
heading("4.2 Descomposicion periodo a periodo", 2)
source('code', "src/lib/simulator.js (calculo de interestUvc, amortUvc, balanceUvc).")
formula([
    "Interes_UVC(t) = Saldo_UVC(t-1) x tasa_anual x (dias_periodo / base)",
    "Amort_UVC(t)   = Cuota_UVC - Interes_UVC(t)",
    "Saldo_UVC(t)   = Saldo_UVC(t-1) - Amort_UVC(t)",
])

# ============================ 5. BASE DE DIAS ============================
heading("5. Base de dias (convencion de computo)", 1)
source('doc', "Modelos Excel ('DIAS 360 / DEVENGO 30' en TABLA CARTERA SIMULADA). La opcion Actual/365 es elaboracion propia.")
table(
    ["Convencion", "Dias del periodo", "Base anual", "Uso"],
    [
        ["30/360", "Cada mes = 30 dias", "360", "Predominante en los modelos UVCC del BCV."],
        ["Actual/365", "Dias calendario reales", "365", "Opcion alternativa (elaboracion propia)."],
    ],
    widths=[1.5, 2.2, 1.0, 2.1],
)
formula(["Tasa_periodo = Tasa_anual x (dias_periodo / base)"])

# ============================ 6. INTERESES EN BS ============================
heading("6. Valoracion de los intereses en bolivares", 1)
source('doc', "Modelos Excel (Nayrobis: 'Interes cap. Base' e 'Interes Var. Capital'); documentacion-tecnica.md.")
para("El interes se acumula en UVC y se convierte a bolivares. Hay dos modos:")
para("IDI al vencimiento: Interes_Bs = Interes_UVC x IDI_vencimiento.", bullet=True)
para("IDI diario (devengo): se suma el interes diario valorado por el IDI de cada dia.", bullet=True)
formula([
    "IDI al vencimiento:  Interes_Bs = Interes_UVC x IDI_venc",
    "IDI diario:          Interes_Bs = SUM_d ( Interes_UVC_dia x IDI_dia )",
])
heading("6.1 Componente base y componente por variacion (actualizacion)", 2)
source('doc', "Modelos Excel (columnas 'Amortizacion cap inicial' vs 'Componente de actualizacion'); cuentas 513.01.M.35 / 513.01.M.36.")
para("Cada monto en Bs se separa en el valor al IDI de desembolso (componente base) y la diferencia por la variacion del IDI (componente de actualizacion):")
formula([
    "Componente_base_Bs = Monto_UVC x IDI_desembolso",
    "Componente_variacion_Bs = Monto_Bs - Componente_base_Bs",
])

# ============================ 7. CUOTA EN BS ============================
heading("7. Desglose de la cuota en bolivares", 1)
source('code', "src/lib/simulator.js (interestBs, amortBs, cuotaBs, balanceBs).")
formula([
    "Amort_Bs = Amort_UVC x IDI_vencimiento",
    "Cuota_Bs = Interes_Bs + Amort_Bs",
    "Saldo_Bs = Saldo_UVC_final x IDI_vencimiento",
])

# ============================ 8. VALORIZACION UVC ============================
heading("8. Valorizacion UVC (ganancia por inflacion)", 1)
source('code', "src/lib/simulator.js (valorUvcCapital, valorUvcRend). Concepto contable: Manual SUDEBAN.")
para("Es la ganancia en bolivares por la variacion del IDI entre el desembolso y el vencimiento de la cuota:")
formula([
    "Val_UVC_capital = Amort_Bs - Amort_UVC x IDI_desembolso",
    "Val_UVC_rend    = Interes_Bs - Interes_UVC x IDI_desembolso",
])
para("Si el IDI sube (inflacion), la valorizacion es positiva: el banco recibe mas bolivares por la misma cantidad de UVC.")

# ============================ 9. MORA ============================
heading("9. Calculo de la mora", 1)
source('code', "src/lib/simulator.js (daysLate, moraUvc, moraBs). Topes: Resolucion 21-01-02 Art. 7.")
formula([
    "Dias_mora = max(0, dias(vencimiento -> pago) - dias_gracia)",
    "Base_mora_UVC = Amort_UVC   (base 'amortizacion')  o  Saldo_UVC (base 'saldo')",
    "Mora_UVC = Base_mora_UVC x tasa_mora x (Dias_mora / base)",
    "Mora_Bs  = Mora_UVC x IDI_fecha_pago   (o suma diaria si IDI diario)",
])
para("La mora se calcula solo sobre el capital vencido, nunca sobre intereses ni cuotas futuras (sin anatocismo).")
heading("9.1 Topes de la tasa de mora", 2)
source('doc', "Resolucion BCV 21-01-02, Art. 7 y Paragrafo Unico.")
table(
    ["Tipo de credito", "Mora maxima anual adicional"],
    [["Expresado en UVC", "0,80%"], ["No expresado en UVC", "3,00%"]],
    widths=[3.5, 2.5],
)

# ============================ 10. CLASIFICACION ============================
heading("10. Clasificacion de la cartera por dias de mora", 1)
source('own', "Los umbrales 30/60/90 y la nomenclatura provienen del codigo del proyecto; la base normativa (Res. SUDEBAN 009.16) NO consta en la carpeta documentacion.")
table(
    ["Clasificacion", "Dias de mora", "Tratamiento"],
    [
        ["AL DIA", "0", "Sin atraso."],
        ["MORA 1", "1 a 30", "Atraso leve, en gestion de cobro."],
        ["VENCIDO", "31 a 60", "Atraso moderado."],
        ["VENCIDO 2", "61 a 90", "Atraso significativo."],
        ["CASTIGO", "> 90", "Incobrable / castigo."],
    ],
    widths=[1.8, 1.5, 2.7],
)
heading("10.1 Cuentas activas vs. cuentas de orden", 2)
source('own', "El uso generico de los grupos 143 (activo) y 819 (orden) es criterio del sistema; el plan documentado usa 131.35/513.01/358 (ver seccion 12).")
para("Mientras el atraso no supera el umbral 'mora2', el interes y la mora se reportan en cartera activa. Al superarlo, se trasladan a cuentas de orden (no se reconocen como ingreso hasta cobrarse).")

# ============================ 11. CONGELAMIENTO ============================
heading("11. Congelamiento del credito vencido", 1)
source('doc', "Minuta reunion SUDEBAN/BCV/ABV del 17/12/2019, punto 4.")
para("A la fecha en que el credito pasa a vencido, se 'congela': deja de actualizarse por el IDI en las cuentas reales (Cartera y Patrimonio). Las revalorizaciones de capital y de rendimientos se devengan en cuentas de orden hasta que el cliente pague; al cobrarse, se registran en resultados y patrimonio.")
para("En el sistema, cuando una cuota supera 'mora2', la valorizacion se enruta a los campos de orden (valorUvcCapitalOrder / valorUvcRendOrder, bandera 'frozen').")

# ============================ 12. PLAN DE CUENTAS ============================
heading("12. Plan de cuentas contable", 1)
source('doc', "Modificacion Manual de Contabilidad SUDEBAN (Circular SIB-II-GGR-GNP-12161); presentacion 'Impacto contable UVCC'.")
table(
    ["Codigo", "Nombre", "Uso"],
    [
        ["131.35", "Creditos comerciales vigentes objeto de las medidas del BCV", "Capital base."],
        ["131.36", "Variacion de creditos comerciales vigentes", "Variacion del capital (M.01 incremento / M.02 disminucion)."],
        ["358.01", "Variacion de creditos comerciales (patrimonio)", "Contrapartida patrimonial de la variacion del capital."],
        ["513.01.M.35", "Rendimientos por creditos comerciales vigentes", "Interes del componente base."],
        ["513.01.M.36", "Rendimientos por variacion de creditos comerciales", "Interes del componente de actualizacion."],
        ["1110", "Banco", "Efectivo (elaboracion propia: codigo generico)."],
        ["138.00", "Rendimientos por cobrar", "Interes y mora por cobrar (elaboracion propia)."],
        ["532.00", "Comisiones flat por desembolso", "Comision flat (<=0,50%); codigo de elaboracion propia."],
    ],
    widths=[1.1, 3.1, 1.8],
)

# ============================ 13. ASIENTOS ============================
heading("13. Asientos contables", 1)
source('code', "src/lib/simulator.js (buildLedger), con el plan de cuentas documentado.")
heading("13.1 Desembolso", 2)
source('doc', "Comision flat: Aviso Oficial G.O. 41.742. Cuentas: Manual SUDEBAN.")
table(["Debe", "Haber", "Descripcion"],
      [["131.35 Creditos comerciales (capital)", "", "Monto desembolsado en Bs"],
       ["", "1110 Banco", "Efectivo entregado al cliente"],
       ["", "532.00 Comision flat", "Comision retenida (max. 0,50%)"]],
      widths=[2.5,2.0,2.0])
heading("13.2 Devengo de intereses (cada periodo)", 2)
source('doc', "Separacion base/variacion segun cuentas 513.01.M.35 y 513.01.M.36.")
table(["Debe", "Haber"],
      [["138.00 Rendimientos por cobrar", ""],
       ["", "513.01.M.35 Rendimiento (componente base)"],
       ["", "513.01.M.36 Rendimiento (componente de actualizacion)"]],
      widths=[3.0,3.5])
heading("13.3 Variacion de capital (actualizacion por IDI)", 2)
source('doc', "Cuentas 131.36 / 358.01 (Manual SUDEBAN; presentacion contable).")
table(["Debe", "Haber"],
      [["131.36 Variacion de creditos comerciales", ""],
       ["", "358.01 Variacion de creditos comerciales (patrimonio)"]],
      widths=[3.0,3.5])
para("Una disminucion del IDI invierte el asiento (debe 358.01 / haber 131.36).")
heading("13.4 Devengo de mora y cobro de la cuota", 2)
source('code', "src/lib/simulator.js (buildLedger).")
table(["Evento", "Debe", "Haber"],
      [["Devengo mora", "138.00 Rendimientos por cobrar (mora)", "513.01.M.35 Ingresos por mora"],
       ["Cobro de cuota", "1110 Banco", "131.35 capital / 138.00 intereses / mora"]],
      widths=[1.4,2.6,2.5])

# ============================ 14. PISO DE IDI ============================
heading("14. Cancelacion anticipada y piso de IDI", 1)
source('doc', "Resolucion BCV 21-01-02, Art. 5 lit. b/c y Art. 6; Circular SUDEBAN 13083.")
para("El deudor puede cancelar anticipadamente sin penalidad. Si el IDI de la fecha de pago resulta menor al IDI de otorgamiento, para determinar el monto a pagar se usa el IDI de otorgamiento:")
formula(["IDI_efectivo = max( IDI_fecha_pago , IDI_desembolso )"])
para("En condiciones normales (IDI creciente) no tiene efecto; solo actua si el IDI baja por debajo del de otorgamiento. Parametro: idiFloorOnPrepay (activo por defecto).")

# ============================ 15. PREPAGOS ============================
heading("15. Prepagos y reconduccion del calendario", 1)
source('own', "El derecho a prepagar sin penalidad es regulatorio; las dos politicas de reconduccion son criterio del sistema.")
table(["Politica", "Efecto"],
      [["Reducir plazo (reduce_term)", "Se mantiene la cuota y se acorta el numero de cuotas."],
       ["Reducir cuota (reduce_installment)", "Se mantiene el plazo y baja la cuota."]],
      widths=[2.6,3.9])
formula(["Nueva_Cuota_UVC = Nuevo_Saldo_UVC x [ i / (1 - (1+i)^(-n_restante)) ]"])

# ============================ 16. TOPES Y COMISIONES ============================
heading("16. Topes regulatorios y validacion del sistema", 1)
source('doc', "Resolucion 21-01-02 (Arts. 2, 3, 7) y Aviso Oficial G.O. 41.742 (comision flat).")
table(["Parametro", "Limite", "Fuente"],
      [["Tasa de interes (UVC comercial/microcredito)", "4% a 10%", "Art. 3"],
       ["Tasa de interes (Cartera Productiva)", "2%", "Art. 2"],
       ["Mora (UVC)", "<= 0,80%", "Art. 7"],
       ["Mora (no UVC)", "<= 3%", "Art. 7 P.U."],
       ["Comision flat de desembolso", "<= 0,50%", "Aviso Oficial G.O. 41.742"]],
      widths=[3.0,1.5,2.0])
para("El sistema evalua estos limites (modulo regulatory.js) y, salvo en modo referencia, bloquea la simulacion cuando se exceden las tasas o la comision.")
heading("16.1 Modo referencia / historico", 2)
source('own', "Mecanismo creado por el asistente para reproducir tablas previas a 2021 (Nayrobis 16%/3%).")
para("Cuando 'allowHistoricalRates' esta activo, el exceso de tasa o mora se muestra como alerta en lugar de bloquear. Al desactivarlo, se exigen estrictamente los topes vigentes.")

# ============================ 17. SITUACIONES ============================
heading("17. Situaciones y condiciones operativas", 1)
source('own', "Casuistica integrada por el asistente a partir de las reglas anteriores.")
table(["Situacion", "Comportamiento esperado"],
      [
       ["Pago al dia (sin atraso)", "No hay mora. Interes y capital se reconocen normalmente; clasificacion AL DIA."],
       ["Atraso dentro del umbral (<= mora2)", "Se cobra mora; interes y mora permanecen en cartera activa."],
       ["Atraso sobre el umbral (> mora2)", "Credito en orden/congelado: revalorizacion y rendimientos van a cuentas de orden."],
       ["Prepago con 'reducir plazo'", "Misma cuota, termina antes."],
       ["Prepago con 'reducir cuota'", "Mismo plazo, cuota menor."],
       ["Cancelacion con IDI a la baja", "Se aplica el piso: se usa el IDI de otorgamiento."],
       ["IDI no publicado (fin de semana/feriado)", "Se usa el ultimo IDI habil; futuro se extrapola."],
       ["Vencimiento en dia no habil", "Se mueve al siguiente dia habil (parametro adjustToBusinessDay)."],
       ["Credito no expresado en UVC", "No se aplica IDI; mora hasta 3%."],
       ["Tasa o mora fuera de tope", "Alerta (modo referencia) o bloqueo (modo estricto)."],
      ],
      widths=[2.4,4.1])

# ============================ 18. DIAS HABILES / TIR ============================
heading("18. Dias habiles, feriados y TIR", 1)
source('own', "Ajuste a dia habil, lista de feriados y calculo de TIR: criterio e implementacion del sistema.")
para("Ajuste a dia habil: si el vencimiento cae en fin de semana o feriado, pasa al siguiente dia habil.", bullet=True)
para("TIR: se calcula por biseccion sobre los flujos; la TIR anual = (1 + TIR_mensual)^12 - 1.", bullet=True)

# ============================ 19. EJEMPLO ============================
heading("19. Ejemplo numerico de referencia (tabla Nayrobis)", 1)
source('doc', "Modelo Excel 'Tabla amortizacion nayrobis bolivar'.")
table(["Dato", "Valor"],
      [["Capital", "160.000,00 Bs"],
       ["Fecha de desembolso", "16/10/2025"],
       ["IDI de desembolso", "0,98495915"],
       ["Capital en UVC", "160.000 / 0,98495915 = 162.443,29 UVC"],
       ["Tasa de interes anual", "16% (modo referencia)"],
       ["Tasa de mora", "3% (modo referencia)"],
       ["Interes base diario", "160.000 x 16% / 360 = 71,11 Bs/dia"],
       ["IDI al 14/11/2025", "1,14827445"]],
      widths=[2.6,3.9])
para("El motor del sistema reproduce estos valores, confirmando la consistencia con el modelo oficial.")

# ============================ 20. PARAMETROS ============================
heading("20. Parametros configurables", 1)
source('code', "src/lib/loanStorage.js (initialParams) y simulator.js.")
table(["Parametro", "Descripcion", "Por defecto"],
      [
       ["principal", "Capital en Bs", "160.000"],
       ["annualRate", "Tasa anual %", "16 (referencia)"],
       ["termMonths", "Plazo en meses", "12"],
       ["moraRate", "Tasa de mora anual %", "3 (referencia)"],
       ["disbursementFeeRate", "Comision flat %", "0,50"],
       ["dayCount", "Base de dias", "30/360"],
       ["interestValuation", "IDI diario o al vencimiento", "IDI diario"],
       ["idiFloorOnPrepay", "Piso de IDI", "Activo"],
       ["allowHistoricalRates", "Modo referencia", "Activo"],
       ["mora1/mora2/mora3", "Umbrales de clasificacion", "30/60/90"],
      ],
      widths=[1.9,2.9,1.7])

# ============================ 21. GLOSARIO ============================
heading("21. Glosario", 1)
source('own', "Definiciones redactadas por el asistente.")
table(["Termino", "Definicion"],
      [["UVC / UVCC", "Unidad de Valor de Credito (Comercial). Unidad de cuenta del credito."],
       ["IDI", "Indice de Inversion. Valor en Bs de una UVC, publicado por el BCV."],
       ["Componente base", "Valor de un monto UVC al IDI de desembolso."],
       ["Componente de actualizacion", "Diferencia por la variacion del IDI."],
       ["Mora", "Interes punitorio por atraso, solo sobre capital vencido."],
       ["Cuenta de orden", "Registro fuera de balance para creditos vencidos/congelados."],
       ["TIR", "Tasa interna de retorno del flujo del credito."]],
      widths=[1.8,4.7])

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run("Nota final: verificar siempre la vigencia de las resoluciones del BCV y SUDEBAN antes de usar en produccion. Las secciones marcadas como 'Elaboracion propia' no tienen respaldo en la documentacion entregada.")
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = GRIS

doc.save(OUT)
print("OK ->", OUT)
