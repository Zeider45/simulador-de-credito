#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera la documentacion tecnica de desarrollo del modulo de credito UVC."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/user/simulador-de-credito/docs/Documentacion-Tecnica-Modulo-Credito.docx"

AZUL = RGBColor(0x10, 0x2A, 0x43)
AZUL2 = RGBColor(0x1F, 0x3A, 0x5F)
GRIS = RGBColor(0x55, 0x55, 0x55)
MORADO = RGBColor(0x3B, 0x10, 0x6E)
VERDE = RGBColor(0x0E, 0x4D, 0x1E)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = AZUL if level==1 else AZUL2
    return h

def P(text, bullet=False, num=False):
    style = "List Bullet" if bullet else ("List Number" if num else None)
    return doc.add_paragraph(text, style=style)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'F2F4F7'); pPr.append(sh)
    run = p.add_run(text)
    run.font.name = "Consolas"; run.font.size = Pt(9.5); run.font.color.rgb = MORADO
    return p

def formula(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"; run.font.size = Pt(10); run.font.color.rgb = VERDE
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(c,"102A43")
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""
            run=cells[i].paragraphs[0].add_run(str(v))
            run.font.size=Pt(9)
            if i==0 and ("(" in str(v) or str(v).islower() or "." in str(v)):
                run.font.name="Consolas"
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ===== PORTADA =====
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("Modulo de Credito UVC\nDocumentacion Tecnica de Desarrollo"); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=AZUL
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=s.add_run("Arquitectura, modelo de datos, algoritmos de calculo, contabilidad,\nvalidaciones, contrato de API y casos limite"); rs.font.size=Pt(12); rs.font.color.rgb=GRIS
v=doc.add_paragraph(); v.alignment=WD_ALIGN_PARAGRAPH.CENTER
rv=v.add_run("Version 1.0  ·  Stack: Next.js (App Router) + React  ·  Motor de calculo en JavaScript puro"); rv.font.size=Pt(10); rv.font.color.rgb=GRIS
doc.add_page_break()

# ===== 1. OBJETIVO =====
H("1. Objetivo y alcance del modulo",1)
P("El modulo de credito calcula y administra creditos comerciales denominados en Unidad de Valor de Credito (UVC). Toma los parametros de un prestamo y produce: (1) un cronograma de pagos (schedule) con desglose diario, (2) un resumen de indicadores, (3) los asientos contables (ledger) y (4) un reporte de cumplimiento normativo.")
P("Caracteristicas funcionales:")
for x in ["Conversion capital<->UVC mediante el Indice de Inversion (IDI) diario.",
          "Amortizacion francesa (cuota fija en UVC) y modalidad de pago al vencimiento.",
          "Valoracion de intereses por IDI diario o IDI al vencimiento, con separacion base/variacion.",
          "Mora diaria configurable, clasificacion de cartera y enrutamiento a cuentas de orden.",
          "Prepagos con reconduccion (reducir plazo o cuota) y piso de IDI en cancelacion.",
          "Validacion de topes regulatorios (tasas, mora, comision) con modo bloqueante o de alerta.",
          "Generacion de asientos contables de partida doble.",
          "Calculo de TIR."]:
    P(x, bullet=True)

# ===== 2. ARQUITECTURA =====
H("2. Arquitectura y componentes",1)
P("El modulo separa el motor de calculo (puro, sin dependencias de framework) de la capa de presentacion y persistencia.")
table(["Componente","Archivo","Responsabilidad"],
    [["Motor de calculo","src/lib/simulator.js","Funcion pura simulateLoan(input) -> resultado. Sin estado ni I/O."],
     ["Reglas regulatorias","src/lib/regulatory.js","Limites (REGULATORY_LIMITS) y evaluateCompliance(). Compartido cliente/servidor."],
     ["API","src/app/api/simulate/route.js","POST: inyecta la serie IDI desde SQLite y delega en simulateLoan."],
     ["Persistencia","src/lib/loanStorage.js","Defaults (initialParams), plan de cuentas y CRUD en localStorage."],
     ["UI - listado","src/app/page.js","Lista de creditos, estadisticas."],
     ["UI - detalle","src/app/creditos/[id]/page.js","Formulario, tabla de pagos, asientos, cumplimiento, exportacion."],
     ["UI - detalle diario","src/app/creditos/[id]/detalle/[index]/page.js","Desglose dia a dia de una cuota."],
     ["Fuente IDI","data/idi_series.db / CSV","Serie historica del IDI publicada por el BCV."]],
    widths=[1.5,2.4,2.7])
P("Flujo de una simulacion:")
code("UI (params) -> POST /api/simulate -> carga serie IDI (SQLite) -> simulateLoan(input)\n   -> { summary, schedule, ledger, accounts, compliance, columnHints } -> UI / localStorage")

# ===== 3. CONCEPTOS =====
H("3. Conceptos de dominio",1)
table(["Termino","Definicion"],
    [["UVC","Unidad de Valor de Credito. Unidad de cuenta del prestamo."],
     ["IDI","Indice de Inversion: valor en Bs de 1 UVC en una fecha. Publicado por el BCV."],
     ["IDI de desembolso","IDI en la fecha de liquidacion; fija el capital en UVC y el componente base."],
     ["Componente base","Valor de un monto UVC al IDI de desembolso."],
     ["Componente de variacion","Diferencia respecto al componente base por el cambio del IDI (actualizacion)."],
     ["Cuota de orden","Estado contable de creditos vencidos: revalorizacion fuera de resultados."]],
    widths=[1.8,4.7])
formula(["Capital_UVC = Capital_Bs / IDI_desembolso",
         "Monto_Bs    = Monto_UVC x IDI_fecha"])

# ===== 4. MODELO DE DATOS: ENTRADA =====
H("4. Modelo de datos: parametros de entrada",1)
P("simulateLoan(input) acepta un objeto con los siguientes campos. Los porcentajes se expresan como numero (16 = 16%).")
table(["Campo","Tipo","Default","Descripcion"],
    [["principal","number","160000","Capital nominal en Bs."],
     ["annualRate","number(%)","16","Tasa de interes anual."],
     ["termMonths","int","12","Numero de cuotas (meses)."],
     ["disbursementDate","date","2025-10-16","Fecha de liquidacion."],
     ["firstDueDate","date","2025-11-14","Vencimiento de la primera cuota."],
     ["simulationDays","int","0","Dias simulados desde el desembolso ('hoy' simulado)."],
     ["idi","number","0.98495915","IDI de desembolso (fallback si no hay serie)."],
     ["disbursementFeeRate","number(%)","0.50","Comision flat de desembolso."],
     ["dayCount","enum","30/360","'30/360' o 'actual/365'."],
     ["interestValuation","enum","idi_daily","'idi_due' o 'idi_daily'."],
     ["idiMissing","enum","linear","Relleno de IDI: 'linear' o 'carry'."],
     ["idiFutureStep","number","0.01","Incremento diario del IDI para fechas futuras."],
     ["moraRate","number(%)","3","Tasa de mora anual."],
     ["graceDays","int","0","Dias de gracia antes de generar mora."],
     ["moraBase","enum","amort","Base de la mora: 'amort' o 'saldo'."],
     ["mora1 / mora2 / mora3","int","30 / 60 / 90","Umbrales de clasificacion (dias)."],
     ["creditUvc","bool","true","Aplica IDI (credito en UVC)."],
     ["idiFloorOnPrepay","bool","true","Piso de IDI en cancelacion anticipada."],
     ["allowHistoricalRates","bool","true","Modo referencia: alerta en vez de bloqueo."],
     ["applyPrepay","bool","true","Aplica el excedente del pago a capital."],
     ["recomputeAfterPrepay","bool","true","Reconduce el calendario tras prepago."],
     ["prepayAction","enum","reduce_term","'reduce_term' o 'reduce_installment'."],
     ["idiSeriesText","string","-","Serie IDI 'YYYY-MM-DD,valor' por linea."],
     ["accounts","object|null","null","Override del plan de cuentas."],
     ["adjustToBusinessDay","bool","true","Mueve vencimientos a dia habil."],
     ["holidays","string[]","[]","Feriados 'YYYY-MM-DD'."],
     ["payments","array","[]","Pagos por cuota: {paymentDate, paymentAmount}."],
     ["paymentMode","enum","libre","'libre' o 'simulacion'."]],
    widths=[1.9,1.0,1.1,2.5])

# ===== 5. MODELO DE DATOS: SALIDA =====
H("5. Modelo de datos: estructura de salida",1)
code('{\n  params,        // parametros normalizados (annualRate y moraRate en %)\n  summary,       // indicadores agregados\n  schedule[],    // filas del cronograma (una por cuota)\n  ledger[],      // asientos contables\n  accounts,      // plan de cuentas efectivo\n  compliance,    // reporte de cumplimiento\n  columnHints    // textos de ayuda por columna\n}')
H("5.1 summary",2)
table(["Campo","Descripcion"],
    [["netReceived","principal - comision flat."],
     ["paymentUvc","Cuota fija inicial en UVC."],
     ["avgCuota","Promedio de cuotaBs."],
     ["totalInterest / totalMora / totalAmort / totalCuota","Acumulados en Bs."],
     ["totalOutstanding","Saldo Bs de la ultima cuota."],
     ["annualIrr","TIR anual (o null)."],
     ["asOfDate","Fecha de corte = desembolso + simulationDays."]],
    widths=[2.6,3.9])
H("5.2 schedule[i] (campos principales)",2)
table(["Campo","Descripcion"],
    [["index, dueDate, periodStart, daysPeriod","Identificacion y dias del periodo."],
     ["startBalanceUvc, interestUvc, amortUvc, paymentUvc","Magnitudes en UVC."],
     ["idiDue, idiTextDue","IDI de vencimiento (valor y texto)."],
     ["interestBs, amortBs, cuotaBs, balanceBs","Magnitudes en Bs."],
     ["interesBaseBs, interesVarBs","Interes: componente base y de variacion."],
     ["amortBaseBs, amortVarBs","Amortizacion: componente base y de variacion."],
     ["paymentDate, paymentAmount","Pago registrado."],
     ["daysLate, moraBs","Dias de atraso e interes moratorio."],
     ["activeMora, activeConv, orderMora, orderConv","Rendimientos por estado (activo/orden)."],
     ["moratorio143, moratorio819","Mora clasificada vigente / en orden."],
     ["valorUvcCapital, valorUvcRend","Valorizacion (capital y rendimiento)."],
     ["valorUvcCapitalActive/Order, ...RendActive/Order","Valorizacion enrutada por estado."],
     ["frozen, idiFloorApplied","Banderas: congelado / piso de IDI aplicado."],
     ["status","Clasificacion (AL DIA / MORA 1 / VENCIDO / ...)."],
     ["balanceUvc, balanceBs","Saldo final."],
     ["paidMora, paidInterest, paidPrincipal, paidExtra","Imputacion del pago."],
     ["dailyBreakdown[], moraBreakdown[]","Desglose dia a dia."],
     ["explain","Tooltips con la formula y valores por columna."]],
    widths=[2.9,3.6])

# ===== 6. RESOLUCION IDI =====
H("6. Resolucion del IDI (createIdiResolver)",1)
P("La serie IDI se parsea a una lista ordenada de {date, idi}. El resolver devuelve el IDI para cualquier fecha aplicando estas reglas, con cache por fecha:")
P("Fin de semana: se usa el ultimo dia habil previo (el IDI no cambia sab/dom).", bullet=True)
P("Fecha exacta en la serie: se devuelve ese valor.", bullet=True)
P("Fecha <= hoy sin dato: se arrastra el ultimo valor conocido previo (carry); opcion lineal disponible.", bullet=True)
P("Fecha > hoy (futuro): extrapolacion = ultimo_idi + idiFutureStep x dias_transcurridos.", bullet=True)
P("Sin serie: se usa el idi de fallback del input.", bullet=True)
P("El resolver expone ademas textFor(date) (texto formateado) y sourceFor(date) -> 'BCV' | 'invented'.")

# ===== 7. FECHAS =====
H("7. Utilidades de fecha y base de dias",1)
P("Todas las fechas se manejan en UTC a medianoche para evitar desfases por zona horaria.")
table(["Funcion","Comportamiento"],
    [["addMonths(d,n)","Suma meses ajustando al ultimo dia valido del mes."],
     ["diffDaysActual(a,b)","Dias calendario reales."],
     ["diffDays30360(a,b)","(y2-y1)*360 + (m2-m1)*30 + (min(30,d2)-min(30,d1))."],
     ["isBusinessDay(d)","Falso en sab/dom o feriado de la lista."],
     ["adjustToNextBusinessDay(d)","Avanza al siguiente dia habil (convencion 'following')."]],
    widths=[2.4,4.1])
formula(["base = (dayCount == '30/360') ? 360 : 365",
         "Tasa_periodo = annualRate x (dias_periodo / base)"])

# ===== 8. AMORTIZACION =====
H("8. Motor de amortizacion",1)
H("8.1 Cuota fija en UVC (sistema frances)",2)
formula(["i = annualRate / 12",
         "Cuota_UVC = Capital_UVC x ( i / (1 - (1+i)^(-n)) )",
         "i == 0:  Cuota_UVC = Capital_UVC / n"])
H("8.2 Bucle por periodo",2)
code("para cada cuota t en 1..termMonths:\n  dueDate      = (t==1) ? firstDueDate : addMonths(firstDueDate, t-1)\n  dueDate      = adjustToBusinessDay ? adjustToNextBusinessDay(dueDate) : dueDate\n  daysPeriod   = calcDays(periodStart, dueDate, dayCount)\n  interestUvc  = startBalanceUvc x annualRate x (daysPeriod / base)\n  amortUvc     = Cuota_UVC - interestUvc\n  si es la ultima cuota: amortUvc = startBalanceUvc (cierra el saldo)\n  amortUvc     = max(0, amortUvc)\n  balanceUvc   = startBalanceUvc - (pago ? capital_pagado_UVC : amortUvc)")
P("Si no hay pago registrado, el saldo avanza por la amortizacion teorica; si hay pago, avanza por el capital efectivamente pagado (soporta pagos parciales y prepagos).")
H("8.3 Reconduccion tras prepago",2)
P("reduce_term: se mantiene la cuota y se recalcula el numero de cuotas restantes:")
formula(["n_restante = ceil( -ln(1 - saldo x i / cuota) / ln(1+i) )"])
P("reduce_installment: se mantiene el plazo y se recalcula la cuota sobre el nuevo saldo.")

# ===== 9. INTERESES =====
H("9. Valoracion de intereses en Bs",1)
formula(["idi_due:    interestBs = interestUvc x idiDue",
         "idi_daily:  interestBs = interestUvc x (SUM_d IDI(d) / daysPeriod)",
         "            // d recorre cada dia del periodo (uvcToBsDaily)"])
P("Separacion base/variacion (para cuentas 513.01.M.35 y 513.01.M.36):")
formula(["interesBaseBs = interestUvc x IDI_desembolso",
         "interesVarBs  = interestBs - interesBaseBs",
         "amortBaseBs   = amortUvc  x IDI_desembolso",
         "amortVarBs    = amortBs   - amortBaseBs"])
P("Valorizacion (ganancia por variacion del IDI):")
formula(["valorUvcCapital = amortBs   - amortUvc  x IDI_desembolso",
         "valorUvcRend    = interestBs - interestUvc x IDI_desembolso"])

# ===== 10. PAGOS =====
H("10. Aplicacion de pagos e imputacion",1)
P("Orden de imputacion del monto pagado (paymentAmount):")
P("1) Mora acumulada + mora del periodo (paidMora).", num=True)
P("2) Intereses acumulados + del periodo (paidInterest).", num=True)
P("3) Capital de la cuota (paidPrincipal).", num=True)
P("4) Excedente a capital si applyPrepay (paidExtra).", num=True)
code("remaining   = max(0, paymentAmount)\npaidMora      = min(remaining, unpaidMora + moraBs);      remaining -= paidMora\npaidInterest  = min(remaining, unpaidInterest + interestBs); remaining -= paidInterest\npaidPrincipal = min(remaining, amortBs);                  remaining -= paidPrincipal\npaidExtra     = applyPrepay ? remaining : 0")
P("paymentMode: en 'simulacion', un pago con fecha posterior al 'hoy' simulado (asOfDate) no se reconoce; en 'libre' los pagos se aplican siempre.")

# ===== 11. PISO IDI =====
H("11. Piso de IDI en cancelacion",1)
P("Al convertir el capital pagado UVC<->Bs, el IDI no puede ser menor al de desembolso:")
formula(["idiPayEffective = (creditUvc && idiFloorOnPrepay) ? max(idiPago, IDI_desembolso) : idiPago",
         "capital_pagado_UVC = (paidPrincipal + paidExtra) / idiPayEffective",
         "idiFloorApplied = (idiPago < IDI_desembolso)"])

# ===== 12. MORA =====
H("12. Mora",1)
formula(["daysLate   = max(0, diffDaysActual(dueDate, paymentDate) - graceDays)",
         "baseMoraUvc = (moraBase=='saldo') ? startBalanceUvc : amortUvc",
         "moraUvc     = baseMoraUvc x moraRate x (daysLate / base)",
         "moraBs      = creditUvc ? (idi_daily ? SUM dia : moraUvc x IDI_pago) : moraUvc"])
P("moraBreakdown detalla la mora diaria desde dueDate+graceDays+1 hasta la fecha de corte/pago. La mora se calcula solo sobre el capital, sin anatocismo.")

# ===== 13. CLASIFICACION =====
H("13. Clasificacion y cuentas de orden",1)
code("status = daysLate<=0   -> 'AL DIA'\n         daysLate<=mora1 -> 'MORA 1'\n         daysLate<=mora2 -> 'VENCIDO'\n         daysLate<=mora3 -> 'VENCIDO 2'\n         else            -> 'CASTIGO'")
P("isOrder = daysLate > mora2. Cuando es true (credito 'congelado'):")
P("Rendimientos y mora se reportan en columnas de orden (orderConv, orderMora; moratorio819).", bullet=True)
P("La valorizacion se enruta a valorUvcCapitalOrder / valorUvcRendOrder; frozen = true.", bullet=True)
P("Si es false, va a las columnas activas (activeConv, activeMora; moratorio143) y valor...Active.", bullet=True)

# ===== 14. CONTABILIDAD =====
H("14. Generacion de asientos (buildLedger)",1)
P("Plan de cuentas por defecto (sobrescribible via input.accounts):")
table(["Clave","Codigo","Nombre"],
    [["bank","1110","Banco"],
     ["loan","131.35","Creditos comerciales vigentes (capital base)"],
     ["loanVariation","131.36","Variacion de creditos comerciales"],
     ["equityVariation","358.01","Variacion de creditos comerciales (patrimonio)"],
     ["interestReceivable","138.00","Rendimientos por cobrar"],
     ["interestIncome","513.01.M.35","Rendimientos por creditos comerciales"],
     ["interestVariationIncome","513.01.M.36","Rendimientos por variacion"],
     ["feeIncome","532.00","Comisiones flat por desembolso"]],
    widths=[2.0,1.2,3.3])
P("Asientos generados:")
P("Desembolso: D 131.35 (nominal) / H 1110 (neto) / H 532.00 (comision).", bullet=True)
P("Devengo interes: D 138.00 / H 513.01.M.35 (base) + H 513.01.M.36 (variacion).", bullet=True)
P("Variacion de capital: D 131.36 / H 358.01 (o invertido si el IDI baja).", bullet=True)
P("Devengo mora: D 138.00 / H 513.01.M.35.", bullet=True)
P("Cobro: D 1110 / H 131.35 (capital), 138.00 (interes), mora.", bullet=True)
P("Cada asiento incluye totalDebit y totalCredit; deben cuadrar (partida doble).")

# ===== 15. VALIDACION =====
H("15. Validacion de cumplimiento (regulatory.js)",1)
table(["Limite","Valor"],
    [["uvcRateMin / uvcRateMax","4% / 10%"],
     ["productiveRate","2%"],
     ["uvcMoraMax","0,80%"],
     ["nonUvcMoraMax","3%"],
     ["flatFeeMax","0,50%"]],
    widths=[3.0,3.5])
P("evaluateCompliance(params) devuelve { checks[], violations, blocking[], compliant }. Cada check: { ok, code, message, ref, level }.")
P("Niveles: 'error' (bloquea), 'warning' (alerta), 'info' (ok). Codigos: TASA_INTERES, TASA_MORA, COMISION_FLAT, PISO_IDI.", bullet=True)
P("allowHistoricalRates: degrada TASA_INTERES y TASA_MORA de 'error' a 'warning' (no bloquean).", bullet=True)
P("La UI deshabilita 'Calcular'/'Pagar' si blocking[] no esta vacio.", bullet=True)

# ===== 16. API =====
H("16. Contrato de API",1)
P("POST /api/simulate")
code('// Request body = objeto de parametros (seccion 4). El servidor inyecta\n// idiSeriesText desde data/idi_series.db (tabla idi_series: date, idi, idi_text).\n// Response 200: el objeto resultado (seccion 5).\n// Response 400: { error, details }')
P("GET /api/idi/list -> { rows: [{date, idi, idi_text}] } para poblar el selector de IDI en la UI.")

# ===== 17. PERSISTENCIA =====
H("17. Persistencia (localStorage)",1)
P("Clave: 'simulador-credito.loans'. Registro de credito:")
code('{\n  id, name,\n  params,        // ver seccion 4\n  accountsText,  // JSON del plan de cuentas (editable)\n  payments[],    // {paymentDate, paymentAmount} por cuota\n  result,        // ultimo resultado de simulateLoan (o null)\n  createdAt, updatedAt\n}')

# ===== 18. CALCULO TIR =====
H("18. Calculo de la TIR",1)
P("Se construye el flujo: [netReceived, -pago_1, -pago_2, ...] y se resuelve la TIR mensual por biseccion (busqueda de raiz del VPN). Si no hay cambio de signo, devuelve null.")
formula(["VPN(r) = SUM_t  CF_t / (1+r)^t",
         "annualIrr = (1 + irr_mensual)^12 - 1"])

# ===== 19. CASOS LIMITE =====
H("19. Casos limite e invariantes",1)
table(["Caso","Manejo"],
    [["termMonths = 0","schedule vacio; cuota = 0."],
     ["annualRate = 0","cuota = capital / n; amortizacion lineal."],
     ["Sin serie IDI","se usa el idi de fallback para todas las fechas."],
     ["IDI baja en cancelacion","piso: se usa IDI de desembolso (idiFloorApplied=true)."],
     ["Pago parcial","imputacion en orden; saldo avanza por capital pagado real."],
     ["Pago > cuota","excedente a capital si applyPrepay; reconduccion si aplica."],
     ["Vencimiento en dia no habil","se mueve al siguiente dia habil; afecta dias e IDI."],
     ["Ultima cuota","amortUvc se fuerza al saldo restante (cierra a 0)."],
     ["amortUvc negativa","se trunca a 0."],
     ["Fechas futuras","IDI extrapolado; sourceFor = 'invented'."]],
    widths=[2.3,4.2])
P("Invariantes esperadas: (a) los asientos cuadran (debe = haber); (b) el saldo UVC converge a 0 en la ultima cuota sin prepagos; (c) en escenario de IDI creciente, el piso no altera resultados.")

# ===== 20. CONFIG Y EXTENSION =====
H("20. Configuracion, pruebas y puntos de extension",1)
P("Defaults en src/lib/loanStorage.js (initialParams) y plan de cuentas en DEFAULT_ACCOUNTS.", bullet=True)
P("Validacion numerica: reproducir los modelos de referencia (p. ej. Capital_UVC = 160000/0.98495915 = 162443.29).", bullet=True)
P("Extension sugerida: provisiones por clasificacion, multiples monedas/indices, exportacion contable a formatos del ente.", bullet=True)
P("El motor es JavaScript puro y testeable de forma aislada (sin Next.js).", bullet=True)

doc.add_paragraph()
f=doc.add_paragraph(); fr=f.add_run("Documento tecnico interno del modulo de credito. Las formulas y estructuras reflejan la implementacion en src/lib/simulator.js y src/lib/regulatory.js.")
fr.italic=True; fr.font.size=Pt(9); fr.font.color.rgb=GRIS

doc.save(OUT)
print("OK ->", OUT)
