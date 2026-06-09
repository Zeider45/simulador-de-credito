#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera una documentación Word COMPLETA y detallada del funcionamiento de los
créditos del proyecto: fórmulas, asientos contables, clasificación, condiciones y
situaciones operativas. Pensada para lectura simple (negocio, contabilidad, desarrollo).

Fundamento: src/lib/simulator.js, src/lib/regulatory.js, src/lib/loanStorage.js,
docs/documentacion-tecnica.md, CALCULOS_TABLA_PAGOS.md y la carpeta documentacion/.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/user/simulador-de-credito/docs/Documentacion-Completa-Creditos.docx"

# ---- Paleta de colores ----
AZUL   = RGBColor(0x1F, 0x3A, 0x5F)
AZUL2  = RGBColor(0x2E, 0x5A, 0x88)
GRIS   = RGBColor(0x55, 0x55, 0x55)
VERDE  = RGBColor(0x1B, 0x5E, 0x20)
NARANJA= RGBColor(0x9A, 0x52, 0x00)
MORADO = RGBColor(0x4A, 0x14, 0x8C)
ROJO   = RGBColor(0x8E, 0x1B, 0x1B)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Estilo base legible
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

# Márgenes algo más anchos para lectura
for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

# ---------------------------------------------------------------- helpers
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = AZUL if level == 1 else AZUL2
    return h

def para(text, bullet=False, num=False):
    style = None
    if bullet:
        style = "List Bullet"
    elif num:
        style = "List Number"
    return doc.add_paragraph(text, style=style)

def bold_lead(lead, rest=""):
    p = doc.add_paragraph()
    r = p.add_run(lead)
    r.bold = True
    if rest:
        p.add_run(rest)
    return p

def formula(lines):
    """Bloque monoespaciado para fórmulas."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = MORADO
    return p

def callout(title, text, color=AZUL, fill="EAF1F8"):
    """Caja resaltada (nota / importante / ejemplo)."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def table(headers, rows, widths=None, fontsize=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(fontsize); run.font.color.rgb = BLANCO
        shade(hdr[i], "1F3A5F")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(fontsize)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def ledger_table(title, rows):
    """Tabla de asiento contable: Cuenta | Debe | Haber."""
    bold_lead(title)
    table(["Cuenta contable", "Debe", "Haber"], rows, widths=[4.2, 1.15, 1.15])

# ================================================================ PORTADA
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Manual de Funcionamiento de los Créditos")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = AZUL
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t2.add_run("Créditos comerciales expresados en Unidad de Valor de Crédito (UVC)")
r.font.size = Pt(13); r.font.color.rgb = AZUL2

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Fórmulas de cálculo  ·  Desglose de la cuota  ·  Mora  ·  Clasificación de cartera\n"
                 "Asientos contables  ·  Cancelación anticipada  ·  Situaciones operativas")
rs.font.size = Pt(11); rs.font.color.rgb = GRIS

doc.add_paragraph()
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
rd = d.add_run("Documentación técnica y funcional del Simulador de Crédito\nVersión 1.0")
rd.font.size = Pt(10); rd.font.color.rgb = GRIS

doc.add_paragraph()
callout("¿Para quién es este documento?",
        "Para negocio, contabilidad, auditoría y desarrollo. Explica de la forma más simple posible "
        "cómo debe funcionar un crédito de principio a fin: desde el desembolso hasta el último "
        "pago, pasando por la mora, la clasificación de riesgo y los asientos contables. Cada "
        "número del simulador puede rastrearse hasta su fórmula y su fundamento.",
        color=VERDE, fill="EAF6EC")

# ---- Cómo leer ----
heading("Cómo leer este documento", 2)
para("Todas las fórmulas se muestran en bloques de color morado y con tipografía monoespaciada. "
     "Los términos clave (UVC, IDI, mora, etc.) se definen en el Glosario al final. "
     "Donde es útil, se incluye un ejemplo numérico resaltado en una caja de color.")
para("Las cifras de los ejemplos toman como base el caso por defecto del simulador (la tabla de "
     "amortización de referencia): capital 160.000 Bs, IDI de desembolso 0,98495915, plazo 12 meses, "
     "tasa 16 % anual y mora 3 % anual.")

doc.add_page_break()

# ================================================================ ÍNDICE (manual)
heading("Contenido", 1)
indice = [
    "1. Qué es un crédito UVC (en palabras simples)",
    "2. Marco regulatorio de referencia",
    "3. Unidad de Valor de Crédito (UVC) e Índice de Inversión (IDI)",
    "4. El Índice de Inversión día a día: cómo se resuelve cada fecha",
    "5. Parámetros que definen un crédito",
    "6. Sistema de amortización francés (cuota fija en UVC)",
    "7. Base de días (30/360 y actual/365)",
    "8. Valoración de intereses: IDI al vencimiento vs. IDI diario",
    "9. Desglose de la cuota en bolívares (base, variación y valorización UVC)",
    "10. Saldo del crédito",
    "11. Mora: cuándo, cuánto y sobre qué se cobra",
    "12. Orden de imputación de los pagos",
    "13. Clasificación de la cartera de créditos",
    "14. Cartera activa (143) y cuentas de orden (819)",
    "15. Congelamiento del crédito vencido",
    "16. Cancelación anticipada y prepagos",
    "17. Asientos contables (paso a paso)",
    "18. Plan de cuentas SUDEBAN utilizado",
    "19. Ajuste a días hábiles y feriados",
    "20. Comisión flat y cumplimiento regulatorio",
    "21. Tasa Interna de Retorno (TIR)",
    "22. Modo libre vs. modo simulación",
    "23. Catálogo de situaciones y escenarios",
    "24. Resumen de todas las fórmulas (hoja de referencia)",
    "25. Glosario",
]
for it in indice:
    p = doc.add_paragraph(it)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ================================================================ 1
heading("1. Qué es un crédito UVC (en palabras simples)", 1)
para("Un crédito UVC es un préstamo en bolívares cuyo valor se «guarda» en una unidad estable "
     "llamada Unidad de Valor de Crédito (UVC), para que la inflación no licúe el capital prestado. "
     "Funciona en tres pasos:")
para("Al desembolsar, el monto en bolívares se convierte a UVC usando el Índice de Inversión (IDI) "
     "del día. Esa cantidad de UVC es la deuda «real».", num=True)
para("Durante la vida del crédito, la deuda se lleva en UVC: los intereses, la amortización y el "
     "saldo se calculan en UVC con una cuota fija (sistema francés).", num=True)
para("Cada vez que hay que cobrar o contabilizar, las UVC se vuelven a expresar en bolívares "
     "multiplicando por el IDI de esa fecha. Como el IDI sube con la inflación, el banco recupera "
     "el poder adquisitivo y el cliente paga el valor real.", num=True)
callout("En una frase:",
        "el crédito «vive» en UVC (valor estable) y solo se «traduce» a bolívares en el momento de "
        "pagar o contabilizar, usando el IDI de cada fecha.")
para("Si el crédito NO es UVC (parámetro creditUvc desactivado), todo el cálculo se hace "
     "directamente en bolívares: el IDI se trata como 1 y desaparecen los componentes de "
     "actualización. El resto del modelo (cuota francesa, mora, clasificación, asientos) es igual.")

# ================================================================ 2
heading("2. Marco regulatorio de referencia", 1)
para("El modelo implementa la normativa venezolana sobre créditos comerciales en UVC. Los "
     "instrumentos principales son:")
table(
    ["Instrumento", "Fecha / Gaceta", "Qué aporta al modelo"],
    [
        ["Resolución BCV N.° 19-09-01", "G.O. 41.742 - 21/10/2019",
         "Crea la obligación de expresar los créditos comerciales en UVC y dividir el monto en Bs entre el IDI del otorgamiento."],
        ["Aviso Oficial BCV", "G.O. 41.742 - 21/10/2019",
         "Comisión flat máxima de 0,50 % del monto del crédito."],
        ["Circular de entrada en vigencia (BCV)", "24/10/2019",
         "Aplica la Resolución 19-09-01 desde el 28/10/2019."],
        ["Circular SUDEBAN SIB-DSB-CJ-OD-13083", "14/11/2019",
         "Cláusulas mínimas del contrato: amortización de capital en UVC, cancelación anticipada con piso de IDI, % de mora."],
        ["Modif. Manual de Contabilidad (SIB-II-GGR-GNP-12161)", "28/10/2019",
         "Crea las subcuentas 131.35, 131.36, 358.01, 513.01.M.35 y 513.01.M.36."],
        ["Minuta SUDEBAN / BCV / ABV", "17/12/2019",
         "Tratamiento del crédito vencido: congelamiento y registro en cuentas de orden / patrimonio (358.00)."],
        ["Resolución BCV N.° 21-01-02 (VIGENTE)", "G.O. 42.050 - 19/01/2021",
         "Norma actual: define UVC, IDI, tasas 4 %-10 % (2 % Cartera Productiva), mora máxima 0,80 %, piso de IDI en cancelación anticipada."],
    ],
    widths=[2.4, 1.6, 2.7],
)
callout("Nota de vigencia:",
        "las resoluciones se actualizan con frecuencia. Antes de usar en producción, verifique las "
        "versiones vigentes en bcv.org.ve y sudeban.gob.ve.", color=NARANJA, fill="FCF3E7")

# ================================================================ 3
heading("3. Unidad de Valor de Crédito (UVC) e Índice de Inversión (IDI)", 1)
para("La UVC es una unidad de cuenta creada por el BCV para preservar el valor real del crédito. "
     "El IDI (Índice de Inversión) es un factor que el BCV publica diariamente y que indica "
     "cuántos bolívares vale una UVC en una fecha determinada. Refleja la variación del tipo de "
     "cambio de referencia de mercado.")

heading("3.1 Conversión del capital al desembolsar", 2)
para("El monto entregado en bolívares se convierte a UVC con el IDI del día del desembolso. Ese "
     "IDI de otorgamiento queda fijo para toda la vida del crédito (es la referencia «base»).")
formula(["Capital_UVC = Capital_Bs / IDI_desembolso"])
callout("Ejemplo:", "160.000 Bs / 0,98495915 = 162.443,29 UVC.  Esa es la deuda real del cliente.")

heading("3.2 Conversión de UVC a bolívares en cualquier fecha", 2)
formula(["Monto_Bs = Monto_UVC x IDI_fecha"])
para("Esta es la regla que se usa para expresar en bolívares la cuota, los intereses, la "
     "amortización y el saldo en cualquier momento.")

# ================================================================ 4
heading("4. El Índice de Inversión día a día: cómo se resuelve cada fecha", 1)
para("El BCV no publica IDI todos los días (no hay fines de semana ni feriados). El sistema "
     "(función createIdiResolver) resuelve el IDI de cualquier fecha así:")
table(
    ["Situación de la fecha", "Qué IDI se usa"],
    [
        ["Hay dato publicado para esa fecha", "Se usa el valor exacto del BCV."],
        ["Fin de semana (sábado/domingo)", "Se usa el IDI del último día hábil previo (el IDI no cambia en fin de semana)."],
        ["Fecha pasada o de hoy sin dato exacto", "Se arrastra el último IDI conocido (o se interpola, según el parámetro idiMissing)."],
        ["Fecha futura (posterior a hoy)", "Se extrapola: último IDI + (paso diario idiFutureStep x días futuros)."],
        ["Antes del primer dato de la serie", "Se usa el primer IDI disponible."],
        ["No hay serie cargada", "Se usa el IDI por defecto (parámetro idi)."],
    ],
    widths=[3.1, 3.6],
)
para("Cada fila del calendario marca el origen del IDI («BCV» cuando deriva de un dato real hasta "
     "hoy, o «invented» cuando es una extrapolación futura), útil para auditoría.")
callout("Por qué importa:",
        "el IDI elegido para cada fecha afecta directamente los bolívares de intereses, "
        "amortización, mora y saldo. Por eso el origen del dato (real vs. estimado) se hace visible.")

# ================================================================ 5
heading("5. Parámetros que definen un crédito", 1)
para("Un crédito se configura con los siguientes parámetros (los valores por defecto reproducen "
     "la tabla de amortización de referencia):")
table(
    ["Parámetro", "Qué significa", "Valor por defecto"],
    [
        ["principal", "Capital desembolsado en bolívares", "160.000"],
        ["annualRate", "Tasa de interés anual (%)", "16"],
        ["termMonths", "Plazo en meses (número de cuotas)", "12"],
        ["disbursementDate", "Fecha de desembolso", "2025-10-16"],
        ["firstDueDate", "Fecha de vencimiento de la 1.ª cuota", "2025-11-14"],
        ["idi", "IDI del día de desembolso (referencia base)", "0,98495915"],
        ["disbursementFeeRate", "Comisión flat de desembolso (%)", "0,50"],
        ["dayCount", "Convención de días: 30/360 o actual/365", "30/360"],
        ["interestValuation", "Valoración del interés: idi_due o idi_daily", "idi_daily"],
        ["idiMissing", "Relleno de IDI faltante: linear o carry", "linear"],
        ["idiFutureStep", "Incremento diario del IDI para fechas futuras", "0,01"],
        ["moraRate", "Tasa de mora anual (%)", "3"],
        ["graceDays", "Días de gracia antes de generar mora", "0"],
        ["moraBase", "Base de la mora: amort (cuota de capital) o saldo", "amort"],
        ["mora1 / mora2 / mora3", "Umbrales de clasificación (días)", "30 / 60 / 90"],
        ["creditUvc", "Activa la denominación en UVC", "true"],
        ["idiFloorOnPrepay", "Piso de IDI en cancelación anticipada", "true"],
        ["allowHistoricalRates", "Modo referencia: admite tasas históricas sin bloquear", "true"],
        ["applyPrepay", "Aplica el exceso de pago a capital (prepago)", "true"],
        ["recomputeAfterPrepay", "Reconduce el calendario tras un prepago", "true"],
        ["prepayAction", "Política de prepago: reduce_term o reduce_installment", "reduce_term"],
        ["adjustToBusinessDay", "Mueve vencimientos a día hábil siguiente", "true"],
        ["holidays", "Lista de feriados (AAAA-MM-DD)", "[ ]"],
        ["paymentMode", "libre (pagos reales) o simulación (con «hoy» simulado)", "libre"],
    ],
    widths=[1.9, 3.5, 1.3], fontsize=9,
)

# ================================================================ 6
heading("6. Sistema de amortización francés (cuota fija en UVC)", 1)
para("La norma exige que cada cuota incluya intereses y una porción de amortización de capital "
     "expresada en UVC. El sistema usa amortización francesa: la cuota en UVC es fija durante todo "
     "el plazo. Es el método más transparente porque el cliente conoce su cuota desde el inicio.")

heading("6.1 Fórmula de la cuota fija", 2)
formula([
    "i = tasa_anual / 12                          (tasa mensual)",
    "Cuota_UVC = Capital_UVC x [ i / (1 - (1 + i)^(-n)) ]",
    "   n = número de cuotas (meses)",
])
bold_lead("Caso especial - tasa cero: ", "si la tasa es 0, la cuota es simplemente el capital "
          "dividido entre el número de cuotas:")
formula(["Cuota_UVC = Capital_UVC / n"])

heading("6.2 Descomposición de cada cuota", 2)
para("En cada período se calcula primero el interés y luego la amortización es el resto de la cuota:")
formula([
    "tasa_periodo = tasa_anual x (dias_periodo / base_dias)",
    "Interes_UVC  = Saldo_UVC_inicial x tasa_periodo",
    "Amort_UVC    = Cuota_UVC - Interes_UVC",
    "Saldo_UVC_fin = Saldo_UVC_inicial - Amort_UVC",
])
para("Propiedad garantizada del método francés: la amortización crece y el interés decrece cada "
     "período, de modo que el saldo en UVC llega a cero en la última cuota.")

heading("6.3 La última cuota cierra el saldo", 2)
para("En el último período, la amortización se fuerza a igualar exactamente el saldo pendiente y "
     "la cuota se recalcula como interés + ese saldo. Así se evita que queden céntimos de UVC sin "
     "amortizar por redondeo.")
formula([
    "Amort_UVC(ultima)  = Saldo_UVC_inicial",
    "Cuota_UVC(ultima)  = Interes_UVC + Saldo_UVC_inicial",
])
callout("Regla de seguridad:", "si por algún cálculo la amortización diera negativa, se fija en 0 "
        "(la cuota nunca aumenta el capital: no hay anatocismo).")

# ================================================================ 7
heading("7. Base de días (30/360 y actual/365)", 1)
para("La «base de días» define cómo se cuentan los días de cada período y el denominador anual. "
     "Afecta el interés y la mora.")
table(
    ["Convención", "Cómo cuenta los días", "Base anual", "Uso típico"],
    [
        ["30/360", "Cada mes = 30 días; año = 360", "360", "Hipotecario, comercial estructurado (predominante en banca venezolana)"],
        ["Actual/365", "Días calendario reales", "365", "Consumo, microcrédito, productos a tasa de mercado"],
    ],
    widths=[1.3, 2.6, 1.0, 1.8],
)
formula(["tasa_periodo = tasa_anual x (dias_periodo / base_dias)"])

# ================================================================ 8
heading("8. Valoración de intereses: IDI al vencimiento vs. IDI diario", 1)
para("Los intereses se acumulan en UVC, pero hay que expresarlos en bolívares. La pregunta es: "
     "qué IDI se usa para esa conversión. Hay dos modalidades.")

heading("8.1 IDI al vencimiento (idi_due)", 2)
para("Sencillo: convierte todo el interés del período con el IDI del día de vencimiento.")
formula(["Interes_Bs = Interes_UVC x IDI_vencimiento"])

heading("8.2 IDI diario (idi_daily) - recomendado", 2)
para("Más preciso: reparte el interés día a día y cada porción se convierte con el IDI de su propio "
     "día. Es el criterio de devengo diario que usan los modelos contables.")
formula([
    "Interes_UVC_dia = Saldo_UVC_inicial x (tasa_anual / base_dias)",
    "Interes_Bs = SUMA sobre cada dia t del periodo de [ Interes_UVC_dia x IDI(t) ]",
    "",
    "Forma equivalente usada en el codigo:",
    "Interes_Bs = Interes_UVC x ( Suma_IDI_del_periodo / dias_periodo )",
])
callout("¿Cuál usar?", "el IDI diario es el más defendible ante una auditoría y el que mejor refleja "
        "la inflación del período. El IDI al vencimiento es una simplificación aceptable para "
        "presentaciones administrativas.")

# ================================================================ 9
heading("9. Desglose de la cuota en bolívares (base, variación y valorización UVC)", 1)
para("Cada cuota en bolívares tiene dos partes: intereses y amortización. A su vez, cada una de "
     "esas partes se separa en un «componente base» (valor al IDI de desembolso) y un «componente "
     "de variación» (la actualización por inflación entre el desembolso y el vencimiento).")

heading("9.1 Componentes de la cuota", 2)
formula([
    "Amort_Bs  = Amort_UVC  x IDI_vencimiento",
    "Cuota_Bs  = Interes_Bs + Amort_Bs",
    "Saldo_Bs  = Saldo_UVC_fin x IDI_vencimiento",
])

heading("9.2 Componente base vs. componente de variación", 2)
formula([
    "Componente_base_Bs      = Monto_UVC x IDI_desembolso",
    "Componente_variacion_Bs = Monto_UVC x (IDI_vencimiento - IDI_desembolso)",
    "                        = Monto_Bs - Componente_base_Bs",
])
para("Esto se calcula por separado para la amortización y para el interés (amortBaseBs / amortVarBs, "
     "interesBaseBs / interesVarBs). Permite reportar la utilidad financiera ordinaria por separado "
     "de la utilidad por actualización monetaria, y alimenta los asientos contables (cuentas .M.35 "
     "vs. .M.36).")

heading("9.3 Valorización UVC", 2)
para("La valorización UVC es la ganancia en bolívares que produce la inflación (subida del IDI) "
     "entre el desembolso y el vencimiento de cada cuota:")
formula([
    "Val_UVC_capital = Amort_Bs   - Amort_UVC   x IDI_desembolso",
    "Val_UVC_rend    = Interes_Bs - Interes_UVC x IDI_desembolso",
])
table(
    ["Situación del IDI", "Resultado de la valorización"],
    [
        ["IDI_vencimiento > IDI_desembolso (inflación)", "Valorización positiva: el banco recibe más Bs por la misma UVC."],
        ["IDI_vencimiento = IDI_desembolso", "Valorización cero: no hubo cambio de valor."],
        ["IDI_vencimiento < IDI_desembolso (deflación del IDI)", "Valorización negativa: se reversa contablemente."],
    ],
    widths=[3.0, 3.7],
)

# ================================================================ 10
heading("10. Saldo del crédito", 1)
para("El saldo se lleva en UVC y se expresa en bolívares cuando se necesita:")
formula([
    "Saldo_UVC_fin = Saldo_UVC_inicial - capital_amortizado_UVC",
    "Saldo_Bs      = Saldo_UVC_fin x IDI_vencimiento",
])
para("Importante: el avance del saldo depende de si hubo pago registrado.")
para("Si hay un pago válido registrado, el saldo baja por el capital realmente pagado (en UVC), "
     "lo que captura pagos parciales y prepagos.", bullet=True)
para("Si todavía no hay pago, el saldo baja por la amortización programada, para que el calendario "
     "muestre el comportamiento normal del método francés aun antes de cobrar.", bullet=True)

# ================================================================ 11
heading("11. Mora: cuándo, cuánto y sobre qué se cobra", 1)
heading("11.1 Días de mora", 2)
para("La mora empieza el primer día de atraso después del vencimiento, descontando los días de "
     "gracia configurados.")
formula(["Dias_mora = max(0, dias_entre(vencimiento, fecha_pago) - dias_gracia)"])

heading("11.2 Fórmula de la mora", 2)
para("La mora se calcula solo sobre el capital vencido (sin anatocismo: nunca sobre intereses ni "
     "sobre cuotas futuras). La base puede ser la cuota de capital o el saldo total:")
formula([
    "Base_mora_UVC = Amort_UVC   (si moraBase = 'amort')   -> opcion por defecto",
    "Base_mora_UVC = Saldo_UVC   (si moraBase = 'saldo')",
    "",
    "Mora_UVC = Base_mora_UVC x tasa_mora x (Dias_mora / base_dias)",
])
para("Conversión a bolívares (igual que el interés, según la modalidad de valoración):")
formula([
    "Mora_Bs = Mora_UVC x IDI_fecha_pago                         (IDI al vencimiento)",
    "Mora_Bs = SUMA diaria [ Mora_UVC_dia x IDI(t) ]             (IDI diario)",
])

heading("11.3 Comparación de bases de mora", 2)
table(
    ["Base", "Sobre qué cobra", "Cuándo se usa"],
    [
        ["amort", "La cuota de capital impagada", "Opción por defecto y de mayor aceptación regulatoria."],
        ["saldo", "El saldo total del crédito", "Más penalizante; solo si el contrato lo especifica."],
    ],
    widths=[1.0, 2.7, 3.0],
)
callout("Tope regulatorio de la mora:",
        "para créditos en UVC la mora máxima es 0,80 % anual adicional; para créditos no UVC, 3 % "
        "anual (Res. BCV 21-01-02, Art. 7). El sistema verifica estos topes automáticamente.",
        color=NARANJA, fill="FCF3E7")

# Ejemplos de mora
heading("11.4 Ejemplos de mora", 2)
callout("Ejemplo A - mora leve (10 días, umbral 30):",
        "I_mora = 13.192,29 x 0,19 x 10 / 360 = 69,62 Bs.  Se considera VIGENTE (no supera el umbral). "
        "La mora se registra en cartera activa (143).", color=VERDE, fill="EAF6EC")
callout("Ejemplo B - mora avanzada (75 días, umbral 30):",
        "I_mora = 13.192,29 x 0,19 x 75 / 360 = 522,17 Bs.  Supera el umbral: los intereses dejan de "
        "reconocerse como ingreso y la mora pasa a cuentas de orden (819).", color=ROJO, fill="FBECEC")

# ================================================================ 12
heading("12. Orden de imputación de los pagos", 1)
para("Cuando entra un pago, se aplica en un orden estricto para que el capital no se reduzca antes "
     "de cubrir los costos del atraso. En el sistema el orden es:")
para("Mora (intereses moratorios pendientes + mora del período).", num=True)
para("Intereses corrientes (vencidos + del período).", num=True)
para("Capital de la cuota (amortización programada).", num=True)
para("Excedente -> abono a capital (prepago), solo si applyPrepay está activo.", num=True)
formula([
    "restante = pago",
    "pagado_mora     = min(restante, mora_pendiente + mora_periodo);   restante -= pagado_mora",
    "pagado_interes  = min(restante, interes_pend + interes_periodo);  restante -= pagado_interes",
    "pagado_capital  = min(restante, amortizacion_Bs);                 restante -= pagado_capital",
    "pagado_extra    = restante   (si applyPrepay)  -> va a capital",
])
callout("Nota legal:", "el orden del marco regulatorio antepone los gastos de cobranza (si "
        "existen) antes de la mora. El sistema no modela gastos de cobranza, por lo que comienza "
        "por la mora.", color=NARANJA, fill="FCF3E7")

# ================================================================ 13
heading("13. Clasificación de la cartera de créditos", 1)
para("Según los días de mora, cada cuota se clasifica en una categoría de riesgo. Los umbrales "
     "(mora1, mora2, mora3) son configurables.")
table(
    ["Clasificación", "Días de mora", "Significado"],
    [
        ["AL DÍA", "0 (o dentro de la gracia)", "Sin atraso. Cartera vigente."],
        ["MORA 1", "1 hasta mora1 (30)", "Atraso leve; en gestión de cobro."],
        ["VENCIDO", "mora1+1 hasta mora2 (60)", "Atraso moderado; alerta temprana."],
        ["VENCIDO 2", "mora2+1 hasta mora3 (90)", "Atraso significativo; provisión específica."],
        ["CASTIGO", "más de mora3 (90)", "Incobrable; se castiga / pasa a orden de cobro."],
    ],
    widths=[1.5, 2.2, 3.0],
)
formula([
    "si Dias_mora <= 0           -> AL DIA",
    "si Dias_mora <= mora1       -> MORA 1",
    "si Dias_mora <= mora2       -> VENCIDO",
    "si Dias_mora <= mora3       -> VENCIDO 2",
    "si Dias_mora >  mora3       -> CASTIGO",
])

# ================================================================ 14
heading("14. Cartera activa (143) y cuentas de orden (819)", 1)
para("El punto de quiebre contable es el umbral mora2. Mientras el atraso no lo supere, la cuota "
     "está «en orden» contablemente y los rendimientos van a cuentas de activo (grupo 143). Cuando "
     "lo supera, los rendimientos y la mora se trasladan a cuentas de orden (grupo 819) y dejan de "
     "reconocerse como ingreso realizado, por prudencia.")
formula([
    "esta_en_orden = Dias_mora > mora2",
    "",
    "si NO esta_en_orden:   Rend conv act = Interes_Bs ;  Rend mora act = Mora_Bs",
    "                       Rend conv ord = 0          ;  Rend mora ord = 0",
    "si SI esta_en_orden:   Rend conv act = 0          ;  Rend mora act = 0",
    "                       Rend conv ord = Interes_Bs ;  Rend mora ord = Mora_Bs",
])
para("Los indicadores «Moratorio 143» y «Moratorio 819» reflejan ÚNICAMENTE el interés moratorio "
     "(la mora), no el rendimiento convencional, que se reporta aparte:")
formula([
    "Moratorio_143 = Rend mora act   (mora en cartera activa, vigente)",
    "Moratorio_819 = Rend mora ord   (mora en cuentas de orden)",
])
table(
    ["Condición", "Cuenta de intereses", "Cuenta de mora"],
    [
        ["Al día / mora leve (<= mora2)", "143.xx Rendimientos por cobrar", "143.xx Mora por cobrar"],
        ["En orden / castigo (> mora2)", "819.xx Rendimientos en orden", "819.xx Mora en orden"],
    ],
    widths=[2.3, 2.3, 2.1],
)

# ================================================================ 15
heading("15. Congelamiento del crédito vencido", 1)
para("Conforme a la minuta SUDEBAN/BCV/ABV del 17/12/2019, cuando una cuota pasa a «en orden» "
     "(supera mora2) el crédito se «congela»:")
para("Deja de someterse a la actualización diaria por IDI en las cuentas reales (Cartera y "
     "Patrimonio).", bullet=True)
para("La revalorización de capital y de rendimientos se devenga en cuentas de orden hasta que el "
     "cliente pague.", bullet=True)
para("Al cobrarse, se reconoce en resultados/patrimonio; el incremento del capital revaluado se "
     "registra en la cuenta patrimonial 358.00 «Variación de Créditos Comerciales».", bullet=True)
para("En el cálculo, la valorización UVC se enruta según la clasificación:")
formula([
    "vigente (activo):  valorUvcCapitalActive , valorUvcRendActive",
    "en orden (frozen): valorUvcCapitalOrder  , valorUvcRendOrder",
])

# ================================================================ 16
heading("16. Cancelación anticipada y prepagos", 1)
para("Cuando el cliente paga más que la cuota exigible, el excedente se abona a capital (si "
     "applyPrepay está activo). El sistema ofrece dos políticas de reconducción del calendario:")
table(
    ["Política", "Qué mantiene fijo", "Efecto para el cliente"],
    [
        ["reduce_term", "La cuota en UVC", "Paga la misma cuota pero termina antes (menos cuotas)."],
        ["reduce_installment", "El plazo restante", "La cuota mensual baja; termina en la misma fecha."],
    ],
    widths=[1.8, 1.9, 3.0],
)
heading("16.1 Recálculo tras el prepago", 2)
para("Reducir plazo: se mantiene la cuota y se recalcula cuántas cuotas faltan.")
formula([
    "n_restante = -ln( 1 - (Saldo_UVC x i) / Cuota_UVC ) / ln(1 + i)",
    "(redondeado hacia arriba; si la cuota ya cubre el saldo, n_restante = 0)",
])
para("Reducir cuota: se mantiene el plazo y se recalcula la cuota sobre el nuevo saldo.")
formula(["Nueva_Cuota_UVC = Saldo_UVC x [ i / (1 - (1 + i)^(-n_restante)) ]"])

heading("16.2 Piso de IDI en cancelación anticipada", 2)
para("La Resolución BCV 21-01-02 (Art. 5 lit. b/c y Art. 6) protege al acreedor: si el IDI del día "
     "de pago es menor que el IDI de otorgamiento, para calcular el monto a pagar se usa el de "
     "otorgamiento. Así el banco recupera al menos el valor nominal en UVC.")
formula(["IDI_efectivo = max( IDI_fecha_pago , IDI_desembolso )"])
callout("¿Cuándo actúa?", "solo cuando el IDI cae por debajo del de otorgamiento (escenario de "
        "deflación del índice). En el escenario normal de IDI creciente, no cambia nada. La fila "
        "marca idiFloorApplied = true cuando el piso se activó.")

# ================================================================ 17
heading("17. Asientos contables (paso a paso)", 1)
para("El sistema genera los asientos (ledger) a partir del calendario. Cada asiento cuadra "
     "(total Debe = total Haber). Se generan en este orden:")

heading("17.1 Asiento de desembolso", 2)
para("Al entregar el crédito. El capital se reconoce en la cartera (131.35); el banco entrega el "
     "neto (capital menos comisión); la comisión flat es un ingreso.")
ledger_table("Desembolso del crédito:", [
    ["131.35 Créditos comerciales (capital base)", "Capital", "—"],
    ["1110 Banco (efectivo entregado)", "—", "Neto = Capital - Comisión"],
    ["532.00 Comisión flat por desembolso (si > 0)", "—", "Comisión"],
    ["2160 Descuento desembolso (si aplica)", "—", "Descuento"],
])

heading("17.2 Devengo de intereses (cada período)", 2)
para("El rendimiento se separa en componente base (513.01.M.35) y componente por variación / "
     "actualización UVC (513.01.M.36). Solo se generan las líneas con monto significativo.")
ledger_table("Devengo interés cuota k:", [
    ["138.00 Rendimientos por cobrar", "Interés_Bs", "—"],
    ["513.01.M.35 Rendimientos por créditos comerciales (base)", "—", "Interés base"],
    ["513.01.M.36 Rendimientos por variación (actualización)", "—", "Interés variación"],
])

heading("17.3 Variación de capital (actualización por IDI)", 2)
para("Registra el cambio de valor del capital amortizado por efecto del IDI. Si la variación es "
     "positiva (IDI sube) se carga 131.36 contra 358.01; si es negativa, se invierte el asiento.")
ledger_table("Variación de capital - incremento (IDI sube):", [
    ["131.36 Variación de créditos comerciales", "Val. capital", "—"],
    ["358.01 Variación de créditos comerciales (patrimonio)", "—", "Val. capital"],
])
ledger_table("Variación de capital - disminución (IDI baja):", [
    ["358.01 Variación de créditos comerciales (patrimonio)", "Val. capital", "—"],
    ["131.36 Variación de créditos comerciales", "—", "Val. capital"],
])

heading("17.4 Devengo de mora (si aplica)", 2)
ledger_table("Devengo mora cuota k:", [
    ["138.00 Rendimientos por cobrar - mora", "Mora_Bs", "—"],
    ["513.01.M.35 Ingresos por mora", "—", "Mora_Bs"],
])

heading("17.5 Cobro de la cuota", 2)
para("Cuando entra el pago, el banco recibe efectivo y se cancelan, en orden, mora, intereses y "
     "capital (incluido el prepago).")
ledger_table("Pago cuota k:", [
    ["1110 Banco", "Pago recibido", "—"],
    ["138.00 Mora por cobrar (si pagó mora)", "—", "Mora pagada"],
    ["138.00 Rendimientos por cobrar (si pagó interés)", "—", "Interés pagado"],
    ["131.35 Créditos comerciales (capital + prepago)", "—", "Capital pagado"],
])

# ================================================================ 18
heading("18. Plan de cuentas SUDEBAN utilizado", 1)
para("Códigos del Manual de Contabilidad para Instituciones Bancarias (modificación del 28/10/2019). "
     "Las cuentas son configurables por crédito.")
table(
    ["Código", "Nombre", "Uso"],
    [
        ["1110", "Banco", "Efectivo entregado / recibido"],
        ["131.35", "Créditos comerciales vigentes objeto de las medidas del BCV", "Capital base de la cartera"],
        ["131.36", "Variación de créditos comerciales vigentes (.M.01 incr / .M.02 dism)", "Actualización del capital por IDI"],
        ["358.01", "Variación de créditos comerciales (patrimonio)", "Contrapartida patrimonial de la variación"],
        ["138.00", "Rendimientos por cobrar por créditos comerciales", "Intereses y mora por cobrar"],
        ["513.01.M.35", "Rendimientos por créditos comerciales vigentes", "Ingreso por interés (componente base)"],
        ["513.01.M.36", "Rendimientos por variación de créditos comerciales", "Ingreso por actualización UVC"],
        ["532.00", "Comisiones flat por desembolso (máx. 0,50 %)", "Ingreso por comisión flat"],
        ["2160", "Descuento desembolso", "Descuento aplicado al desembolso"],
    ],
    widths=[1.2, 3.6, 1.9], fontsize=9,
)

# ================================================================ 19
heading("19. Ajuste a días hábiles y feriados", 1)
para("Si adjustToBusinessDay está activo y un vencimiento cae en fin de semana o feriado, se mueve "
     "al siguiente día hábil (convención «following»). Esto afecta:")
para("Los días del período (en 30/360 puede dar más de 30 días si hay ajuste).", bullet=True)
para("El IDI aplicado (se usa el del día hábil ajustado).", bullet=True)
para("Los feriados se cargan en una lista configurable (formato AAAA-MM-DD). Se recomienda "
     "mantenerla al día con el calendario de días no hábiles bancarios del BCV/SUDEBAN.")

# ================================================================ 20
heading("20. Comisión flat y cumplimiento regulatorio", 1)
para("La comisión flat de desembolso no puede exceder 0,50 % del monto del crédito. El sistema "
     "evalúa el cumplimiento (función evaluateCompliance) y clasifica cada verificación por nivel.")
table(
    ["Verificación", "Límite", "Nivel si falla"],
    [
        ["Tasa de interés (UVC)", "4 % a 10 % (o 2 % Cartera Productiva)", "error (bloquea)"],
        ["Tasa de mora (UVC)", "<= 0,80 % anual adicional", "error (bloquea)"],
        ["Tasa de mora (no UVC)", "<= 3 % anual", "error (bloquea)"],
        ["Comisión flat", "<= 0,50 % del monto", "error (bloquea)"],
        ["Piso de IDI activo", "debe estar activo", "warning (no bloquea)"],
    ],
    widths=[2.3, 2.9, 1.5],
)
callout("Modo referencia / histórico (allowHistoricalRates):",
        "cuando está activo, los excesos de tasa de interés y de mora se reportan como ALERTA en "
        "vez de bloquear. Sirve para reproducir tablas previas a la Resolución 21-01-02 (por "
        "ejemplo, la tabla de referencia con 16 % de interés y 3 % de mora). Al desactivarlo se "
        "exigen estrictamente los topes vigentes.", color=NARANJA, fill="FCF3E7")

# ================================================================ 21
heading("21. Tasa Interna de Retorno (TIR)", 1)
para("La TIR mide el rendimiento efectivo real del crédito para la institución. Se calcula con los "
     "flujos de caja: el neto recibido al inicio (negativo para el banco como salida de efectivo) y "
     "los pagos del cliente como entradas. Se resuelve por bisección buscando la tasa que hace el "
     "Valor Presente Neto igual a cero, y luego se anualiza:")
formula([
    "VPN(r) = SUMA [ flujo_i / (1 + r)^i ] = 0   ->  r mensual (biseccion)",
    "TIR_anual = (1 + r)^12 - 1",
])
para("Si no hay al menos un flujo positivo y uno negativo, la TIR no está definida (devuelve nulo).")

# ================================================================ 22
heading("22. Modo libre vs. modo simulación", 1)
table(
    ["Modo", "Cómo trata los pagos"],
    [
        ["libre", "Registro de pagos reales: todos los pagos se aplican siempre, sin importar su fecha."],
        ["simulación", "Avanza un «hoy» simulado (asOf = desembolso + simulationDays). Los pagos con fecha posterior a ese «hoy» aún no se reconocen."],
    ],
    widths=[1.4, 5.3],
)
para("El «hoy» simulado (asOfDate) también define hasta dónde se genera el desglose diario cuando "
     "no hay pago registrado.")

# ================================================================ 23
heading("23. Catálogo de situaciones y escenarios", 1)
para("Resumen de cómo responde el modelo ante distintas situaciones:")
table(
    ["Situación", "Comportamiento del modelo"],
    [
        ["Pago puntual y completo", "No hay mora. Interés + amortización del período; el saldo baja según lo pagado."],
        ["Pago anticipado (antes del vencimiento)", "Se marca paidEarly. Si hay excedente, se abona a capital (prepago)."],
        ["Pago parcial", "Se imputa por orden (mora, interés, capital). El saldo solo baja por el capital efectivamente pagado; el resto queda pendiente."],
        ["Atraso dentro de la gracia", "No genera mora (días_mora = 0)."],
        ["Atraso 1 a mora2 días", "Genera mora; clasificación MORA 1 / VENCIDO; rendimientos en cartera activa (143)."],
        ["Atraso mayor a mora2", "Crédito «en orden»: rendimientos y mora a cuentas de orden (819); capital congelado; valorización a cuentas de orden."],
        ["Atraso mayor a mora3", "Clasificación CASTIGO (incobrable)."],
        ["Prepago con reduce_term", "Misma cuota, menos cuotas: se acorta el plazo."],
        ["Prepago con reduce_installment", "Mismo plazo, cuota menor."],
        ["IDI sube (inflación normal)", "Valorización positiva; más bolívares por la misma UVC; el piso de IDI no actúa."],
        ["IDI baja por debajo del de otorgamiento", "Valorización negativa (se reversa); el piso de IDI fija el IDI de otorgamiento en cancelaciones."],
        ["Tasa cero", "Cuota = capital / n; sin intereses."],
        ["Vencimiento en día no hábil", "Si el ajuste está activo, se mueve al siguiente día hábil y se recalculan días e IDI."],
        ["Crédito no UVC", "IDI = 1; sin componentes de actualización ni valorización; cálculo directo en Bs."],
        ["Última cuota", "La amortización iguala el saldo restante; el saldo cierra en cero."],
        ["Tasa/mora fuera de tope", "Se bloquea (error) salvo en modo referencia, donde se reporta como alerta."],
    ],
    widths=[2.3, 4.4], fontsize=9,
)

# ================================================================ 24
heading("24. Resumen de todas las fórmulas (hoja de referencia)", 1)
formula([
    "--- Conversion UVC / IDI ---",
    "Capital_UVC = Capital_Bs / IDI_desembolso",
    "Monto_Bs    = Monto_UVC x IDI_fecha",
    "",
    "--- Cuota francesa ---",
    "i          = tasa_anual / 12",
    "Cuota_UVC  = Capital_UVC x [ i / (1 - (1+i)^(-n)) ]      (tasa 0: Capital_UVC / n)",
    "",
    "--- Descomposicion por periodo ---",
    "tasa_periodo  = tasa_anual x (dias_periodo / base_dias)",
    "Interes_UVC   = Saldo_UVC_inicial x tasa_periodo",
    "Amort_UVC     = Cuota_UVC - Interes_UVC      (ultima cuota: Amort_UVC = Saldo_UVC)",
    "Saldo_UVC_fin = Saldo_UVC_inicial - Amort_UVC",
    "",
    "--- Valoracion en bolivares ---",
    "Interes_Bs (idi_due)   = Interes_UVC x IDI_vencimiento",
    "Interes_Bs (idi_daily) = Interes_UVC x (Suma_IDI_periodo / dias_periodo)",
    "Amort_Bs               = Amort_UVC x IDI_vencimiento",
    "Cuota_Bs               = Interes_Bs + Amort_Bs",
    "Saldo_Bs               = Saldo_UVC_fin x IDI_vencimiento",
    "",
    "--- Base / variacion / valorizacion ---",
    "Componente_base_Bs = Monto_UVC x IDI_desembolso",
    "Val_UVC_capital    = Amort_Bs   - Amort_UVC   x IDI_desembolso",
    "Val_UVC_rend       = Interes_Bs - Interes_UVC x IDI_desembolso",
    "",
    "--- Mora ---",
    "Dias_mora = max(0, dias(vencimiento, pago) - dias_gracia)",
    "Mora_UVC  = Base_mora_UVC x tasa_mora x (Dias_mora / base_dias)",
    "Mora_Bs   = Mora_UVC x IDI   (o suma diaria con IDI(t))",
    "",
    "--- Clasificacion ---",
    "AL DIA / MORA 1 / VENCIDO / VENCIDO 2 / CASTIGO  segun mora1, mora2, mora3",
    "en_orden = Dias_mora > mora2   ->  rendimientos y mora a cuentas 819",
    "",
    "--- Prepago / piso IDI ---",
    "IDI_efectivo = max(IDI_pago, IDI_desembolso)",
    "n_restante   = -ln(1 - (Saldo_UVC x i)/Cuota_UVC) / ln(1+i)   (reduce_term)",
    "",
    "--- TIR ---",
    "VPN(r)=SUMA flujo_i/(1+r)^i = 0 ;  TIR_anual = (1+r)^12 - 1",
])

# ================================================================ 25
heading("25. Glosario", 1)
table(
    ["Término", "Definición"],
    [
        ["UVC", "Unidad de Valor de Crédito. Unidad de cuenta del BCV para créditos en moneda nacional, que preserva el valor real del capital."],
        ["IDI", "Índice de Inversión. Factor que el BCV publica a diario; indica cuántos bolívares vale una UVC en una fecha."],
        ["IDI de desembolso / base", "IDI del día del desembolso; referencia fija para toda la vida del crédito."],
        ["Amortización", "Porción de la cuota que abona al capital."],
        ["Cuota", "Pago periódico = intereses + amortización."],
        ["Interés corriente / convencional", "Interés ordinario por el uso del dinero."],
        ["Mora", "Interés punitorio por atraso, sobre el capital vencido (sin anatocismo)."],
        ["Días de gracia", "Días de atraso que no generan mora."],
        ["Base de días", "Convención de conteo (30/360 o actual/365) que define el denominador anual."],
        ["Componente base", "Parte de la cuota valorada al IDI de desembolso."],
        ["Componente de variación", "Parte por actualización (diferencia entre IDI de vencimiento y de desembolso)."],
        ["Valorización UVC", "Ganancia en Bs por la subida del IDI entre desembolso y vencimiento."],
        ["Cartera activa (143)", "Créditos con atraso <= mora2; rendimientos en cuentas de activo."],
        ["Cuentas de orden (819)", "Rendimientos y mora de créditos con atraso > mora2; fuera de resultados hasta cobrar."],
        ["Congelamiento", "El crédito vencido deja de actualizarse por IDI en cuentas reales."],
        ["Piso de IDI", "En cancelación anticipada, el IDI no puede ser menor al de otorgamiento."],
        ["Prepago", "Pago de capital por encima de la cuota exigible."],
        ["TIR", "Tasa Interna de Retorno; rendimiento efectivo del crédito para la institución."],
        ["asOf / hoy simulado", "Fecha de corte en modo simulación (desembolso + simulationDays)."],
        ["SUDEBAN", "Superintendencia de las Instituciones del Sector Bancario."],
        ["BCV", "Banco Central de Venezuela."],
    ],
    widths=[1.9, 4.8], fontsize=9.5,
)

doc.add_paragraph()
fin = doc.add_paragraph()
rf = fin.add_run("Documento de uso interno. Verifique la vigencia de las resoluciones del BCV y "
                 "SUDEBAN antes de implementar en producción. Las fórmulas y reglas aquí descritas "
                 "corresponden al motor de cálculo del proyecto (src/lib/simulator.js, "
                 "regulatory.js y loanStorage.js).")
rf.italic = True; rf.font.size = Pt(9); rf.font.color.rgb = GRIS

doc.save(OUT)
print("OK ->", OUT)
