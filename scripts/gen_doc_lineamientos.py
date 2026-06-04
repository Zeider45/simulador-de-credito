#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera lineamientos y bases para el desarrollo de un modulo de credito UVC.
Enfoque independiente de cualquier implementacion concreta."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/user/simulador-de-credito/docs/Lineamientos-Desarrollo-Modulo-Credito.docx"

AZUL = RGBColor(0x10, 0x2A, 0x43)
AZUL2 = RGBColor(0x1F, 0x3A, 0x5F)
GRIS = RGBColor(0x55, 0x55, 0x55)
VERDE = RGBColor(0x0E, 0x4D, 0x1E)

doc = Document()
n = doc.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11)
n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.15

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def H(text,level=1):
    h=doc.add_heading(text,level=level)
    for r in h.runs: r.font.color.rgb = AZUL if level==1 else AZUL2
    return h

def P(text,bullet=False,num=False):
    return doc.add_paragraph(text, style="List Bullet" if bullet else ("List Number" if num else None))

def G(text):
    """Parrafo de lineamiento destacado."""
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    r=p.add_run("Lineamiento. "); r.bold=True; r.font.color.rgb=VERDE; r.font.size=Pt(11)
    r2=p.add_run(text); r2.font.size=Pt(11)
    return p

def formula(lines):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.space_after=Pt(8)
    r=p.add_run("\n".join(lines)); r.font.name="Consolas"; r.font.size=Pt(10); r.font.color.rgb=VERDE
    return p

def code(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25)
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(8)
    pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'F2F4F7'); pPr.append(sh)
    r=p.add_run(text); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=AZUL2
    return p

def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c,"102A43")
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""; rr=cells[i].paragraphs[0].add_run(str(v)); rr.font.size=Pt(9)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ===== PORTADA =====
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("Lineamientos y Bases para el Desarrollo\nde un Modulo de Credito UVC"); r.bold=True; r.font.size=Pt(23); r.font.color.rgb=AZUL
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=s.add_run("Guia tecnica independiente de la implementacion.\nDefine que debe calcular, como y bajo que reglas un modulo de creditos\nexpresados en Unidad de Valor de Credito (UVC)."); rs.font.size=Pt(12); rs.font.color.rgb=GRIS
v=doc.add_paragraph(); v.alignment=WD_ALIGN_PARAGRAPH.CENTER
rv=v.add_run("Version 1.0  ·  Documento de diseno / requisitos"); rv.font.size=Pt(10); rv.font.color.rgb=GRIS
doc.add_page_break()

# ===== 1 =====
H("1. Proposito y alcance",1)
P("Este documento establece los lineamientos y fundamentos para desarrollar un modulo de credito cuyo capital se expresa en Unidad de Valor de Credito (UVC) y se actualiza mediante un indice diario (Indice de Inversion, IDI). No describe una implementacion concreta: define el comportamiento esperado, las formulas, las reglas de negocio y las decisiones de diseno que cualquier desarrollo del modulo debe respetar.")
P("El modulo debe permitir, a partir de los datos de un prestamo, producir de forma determinista:")
for x in ["Un cronograma de pagos con desglose por periodo (y, deseablemente, por dia).",
          "Indicadores agregados del credito (totales, saldo, rendimiento).",
          "Los asientos contables de partida doble asociados.",
          "Un dictamen de cumplimiento de los limites regulatorios."]:
    P(x, bullet=True)

# ===== 2 =====
H("2. Principios de diseno",1)
G("Separar el motor de calculo (logica pura, sin entrada/salida ni interfaz) de las capas de reglas, servicio (API), persistencia y presentacion. El motor debe ser una funcion determinista: mismas entradas, misma salida.")
G("El calculo no debe tener efectos secundarios ni depender del reloj salvo en un unico punto explicito (la 'fecha de corte' o 'hoy'), que debe ser inyectable para poder reproducir escenarios.")
G("Toda magnitud monetaria debe poder rastrearse hasta su formula y sus insumos (trazabilidad/auditoria). Conviene exponer, por cada valor, la formula y los valores usados.")
G("Los parametros regulatorios y de producto deben ser configurables (no fijos en el codigo), con valores por defecto razonables y validados.")
G("Manejar las fechas en una zona horaria unica y estable (preferiblemente UTC a medianoche) para evitar desfases.")
G("Cuidar la precision numerica: definir el redondeo de presentacion (2 decimales en Bs) por separado del calculo interno (mayor precision en UVC).")

# ===== 3 =====
H("3. Conceptos de dominio",1)
table(["Concepto","Definicion"],
    [["UVC","Unidad de cuenta del credito. El capital se expresa y amortiza en UVC."],
     ["IDI","Indice de Inversion: cuantos bolivares vale 1 UVC en una fecha. Se publica a diario."],
     ["IDI de desembolso","IDI de la fecha de liquidacion. Fija el capital en UVC y el 'componente base'."],
     ["Componente base","Valor de un monto UVC al IDI de desembolso (la parte no afectada por la variacion)."],
     ["Componente de variacion","Diferencia por el cambio del IDI entre desembolso y la fecha (actualizacion)."],
     ["Cartera en orden","Estado del credito vencido cuya actualizacion deja de reconocerse en resultados."]],
    widths=[1.9,4.6])
G("Convertir el capital a UVC al desembolsar y reconvertir a bolivares en cada fecha relevante:")
formula(["Capital_UVC = Capital_Bs / IDI_desembolso",
         "Monto_Bs    = Monto_UVC x IDI_fecha"])

# ===== 4 =====
H("4. Datos de entrada que el modulo debe aceptar",1)
P("El modulo debe permitir configurar al menos los siguientes parametros (los valores indicados son referencias razonables, no obligatorios):")
table(["Parametro","Descripcion","Referencia"],
    [["Capital","Monto desembolsado en Bs","-"],
     ["Tasa de interes anual","Sobre el saldo en UVC","4% a 10%"],
     ["Plazo","Numero de cuotas","-"],
     ["Fecha de desembolso / primera cuota","Fechas base del cronograma","-"],
     ["Comision flat","Cargo unico de desembolso","<= 0,50%"],
     ["Tasa de mora anual","Recargo por atraso","<= 0,80% (UVC)"],
     ["Dias de gracia","Atraso tolerado sin mora","Politica interna"],
     ["Base de dias","30/360 o real/365","30/360 usual"],
     ["Modo de valoracion","IDI diario o IDI al vencimiento","Diario (mas preciso)"],
     ["Umbrales de clasificacion","Dias para cada estado de mora","p. ej. 30/60/90"],
     ["Serie IDI","Valores historicos del indice","Del ente emisor"],
     ["Calendario de feriados","Dias no habiles bancarios","Del ente emisor"],
     ["Pagos","Fecha y monto por cuota","-"]],
    widths=[2.2,2.6,1.7])

# ===== 5 =====
H("5. Datos de salida que el modulo debe producir",1)
P("Cronograma (una fila por cuota), que como minimo incluya:")
for x in ["Saldo, interes y amortizacion en UVC y en Bs.",
          "IDI aplicado al vencimiento y separacion base / variacion de interes y amortizacion.",
          "Dias de atraso, mora, clasificacion y banderas de estado (vigente / en orden).",
          "Valorizacion (ganancia por variacion del IDI) y saldo final."]:
    P(x, bullet=True)
P("Resumen: neto recibido, cuota, totales (interes, mora, amortizacion), saldo pendiente y TIR.")
P("Asientos contables y dictamen de cumplimiento (ver secciones 13 y 14).")
G("Es recomendable producir tambien un desglose diario (interes y mora devengados por dia) para auditoria y para la modalidad de valoracion por IDI diario.")

# ===== 6 =====
H("6. Manejo del indice (IDI)",1)
P("El indice se publica solo en dias habiles. El modulo debe resolver el IDI para cualquier fecha con reglas explicitas y deterministas:")
table(["Situacion","Regla recomendada"],
    [["Fecha con valor publicado","Usar el valor exacto."],
     ["Fin de semana / feriado","Usar el ultimo valor habil previo (el IDI no cambia)."],
     ["Fecha pasada sin dato","Arrastrar el ultimo valor conocido (o interpolacion lineal)."],
     ["Fecha futura","Extrapolar de forma transparente y marcarla como estimada."],
     ["Sin serie disponible","Usar un valor de respaldo declarado."]],
    widths=[2.4,4.1])
G("Distinguir y senalar el origen de cada IDI (publicado vs estimado) para que la auditoria sepa que valores son oficiales.")

# ===== 7 =====
H("7. Fechas y base de dias",1)
G("Si una cuota vence en dia no habil, moverla al siguiente dia habil (convencion 'following'); esto afecta el conteo de dias y el IDI aplicado.")
table(["Base","Conteo de dias del periodo","Base anual"],
    [["30/360","Cada mes cuenta como 30 dias","360"],
     ["Real/365","Dias calendario reales","365"]],
    widths=[1.5,3.5,1.5])
formula(["base = (convencion 30/360) ? 360 : 365",
         "Tasa_periodo = Tasa_anual x (dias_periodo / base)"])

# ===== 8 =====
H("8. Amortizacion",1)
P("El modulo debe soportar, como minimo, dos modalidades: pago en cuotas y pago unico al vencimiento. En ambas, cada pago debe contemplar interes y la porcion de capital correspondiente, expresados en UVC.")
H("8.1 Cuota fija en UVC (sistema frances)",2)
formula(["i = Tasa_anual / 12",
         "Cuota_UVC = Capital_UVC x ( i / (1 - (1+i)^(-n)) )",
         "Si i = 0:  Cuota_UVC = Capital_UVC / n"])
H("8.2 Descomposicion por periodo",2)
formula(["Interes_UVC(t) = Saldo_UVC(t-1) x Tasa_anual x (dias / base)",
         "Amort_UVC(t)   = Cuota_UVC - Interes_UVC(t)",
         "Saldo_UVC(t)   = Saldo_UVC(t-1) - Amort_UVC(t)"])
G("La ultima cuota debe forzar la amortizacion al saldo restante para cerrar el credito exactamente en cero. La amortizacion nunca debe ser negativa.")
H("8.3 Reconduccion tras prepago",2)
P("Ante un pago superior a la cuota, el excedente abona capital y debe ofrecerse reconducir el calendario de dos formas:")
P("Reducir plazo: se mantiene la cuota y se recalculan las cuotas restantes.", bullet=True)
P("Reducir cuota: se mantiene el plazo y se recalcula la cuota sobre el nuevo saldo.", bullet=True)

# ===== 9 =====
H("9. Valoracion de intereses",1)
P("El interes se devenga en UVC y se convierte a bolivares segun la modalidad elegida:")
formula(["IDI al vencimiento:  Interes_Bs = Interes_UVC x IDI_vencimiento",
         "IDI diario:          Interes_Bs = SUM_dia ( Interes_UVC_dia x IDI_dia )"])
G("Separar cada monto en componente base (al IDI de desembolso) y componente de variacion (el resto), porque la contabilidad los registra en cuentas distintas:")
formula(["Componente_base     = Monto_UVC x IDI_desembolso",
         "Componente_variacion = Monto_Bs - Componente_base"])
P("La valorizacion (ganancia por inflacion) es la parte de variacion del capital y del rendimiento.")

# ===== 10 =====
H("10. Aplicacion e imputacion de pagos",1)
P("Cuando se recibe un pago, debe imputarse en un orden definido y trazable:")
P("Mora acumulada y del periodo.", num=True)
P("Intereses acumulados y del periodo.", num=True)
P("Capital de la cuota.", num=True)
P("Excedente a capital (prepago), si esta habilitado.", num=True)
G("El saldo debe avanzar por el capital efectivamente pagado cuando hay pago (soportando pagos parciales), o por la amortizacion teorica cuando aun no hay pago.")
G("Distinguir un modo 'registro real' (los pagos siempre se aplican) de un modo 'simulacion' (un pago con fecha posterior al 'hoy' simulado no se reconoce todavia).")

# ===== 11 =====
H("11. Cancelacion anticipada y piso del indice",1)
P("El deudor puede cancelar anticipadamente sin penalidad. Para proteger el valor del credito:")
G("Si el IDI de la fecha de cancelacion es menor al IDI de desembolso, para determinar el monto a pagar debe usarse el IDI de desembolso (piso).")
formula(["IDI_efectivo = max( IDI_fecha_pago , IDI_desembolso )"])
P("En escenarios de IDI creciente el piso no tiene efecto; solo actua cuando el indice cae por debajo del de desembolso.")

# ===== 12 =====
H("12. Mora",1)
formula(["Dias_mora = max(0, dias(vencimiento -> pago) - dias_gracia)",
         "Base_mora = capital de la cuota (o saldo, segun politica)",
         "Mora_UVC  = Base_mora x Tasa_mora x (Dias_mora / base)",
         "Mora_Bs   = Mora_UVC x IDI (al pago o sumado por dia)"])
G("La mora se calcula solo sobre el capital vencido, nunca sobre intereses ni cuotas futuras (sin anatocismo), y debe respetar el tope regulatorio.")

# ===== 13 =====
H("13. Clasificacion de cartera y cuentas de orden",1)
P("El credito se clasifica por dias de atraso. Un esquema de referencia:")
table(["Estado","Dias de atraso"],
    [["Al dia","0"],["Mora leve","1 a primer umbral"],
     ["Vencido","hasta segundo umbral"],["Vencido avanzado","hasta tercer umbral"],
     ["Castigo","mayor al tercer umbral"]],
    widths=[2.5,2.5])
G("Al superar el umbral de 'vencido', el credito se 'congela': deja de actualizarse por el IDI en las cuentas reales y su revalorizacion y rendimientos se devengan en cuentas de orden hasta el cobro efectivo. Los umbrales deben ser configurables segun la normativa vigente del ente supervisor.")

# ===== 14 =====
H("14. Contabilidad y asientos",1)
P("El modulo debe generar asientos de partida doble que cuadren (debe = haber). Conceptos minimos a registrar:")
table(["Evento","Debe","Haber"],
    [["Desembolso","Cartera (capital)","Banco (neto) y Comision flat"],
     ["Devengo de interes","Rendimientos por cobrar","Ingreso base + Ingreso por variacion"],
     ["Variacion de capital","Variacion de cartera (activo)","Variacion de capital (patrimonio)"],
     ["Devengo de mora","Rendimientos por cobrar (mora)","Ingreso por mora"],
     ["Cobro de cuota","Banco","Cartera, intereses y mora"]],
    widths=[1.8,2.3,2.4])
G("Separar el rendimiento ordinario (componente base) del rendimiento por actualizacion (componente de variacion) en cuentas distintas. El incremento del capital por el indice se lleva contra una cuenta patrimonial hasta su cobro.")
G("El plan de cuentas debe ser parametrizable para adaptarse al manual contable vigente del ente supervisor.")

# ===== 15 =====
H("15. Validacion de cumplimiento",1)
P("El modulo debe verificar los limites regulatorios y emitir un dictamen. Valores de referencia frecuentes:")
table(["Limite","Valor de referencia"],
    [["Tasa de interes (UVC)","4% a 10% (2% en carteras dirigidas)"],
     ["Tasa de mora (UVC)","<= 0,80%"],
     ["Tasa de mora (no UVC)","<= 3%"],
     ["Comision flat","<= 0,50%"]],
    widths=[3.0,3.5])
G("Clasificar cada verificacion por severidad (bloqueante / alerta / informativa). Debe poder operar en modo estricto (impide guardar/calcular fuera de los limites) o en modo permisivo controlado (alerta) para reproducir escenarios historicos.")
G("Los limites deben residir en una capa de reglas configurable, separada del motor de calculo, y citar su fundamento normativo.")

# ===== 16 =====
H("16. Interfaz de servicio (contrato recomendado)",1)
P("Se recomienda exponer el calculo como un servicio sin estado que reciba los parametros y devuelva el resultado completo:")
code("Entrada:  { parametros del credito, serie IDI, pagos, plan de cuentas }\nSalida:   { resumen, cronograma, asientos, cumplimiento }\nErrores:  estructura explicita { error, detalle } con codigo de estado")
G("La serie IDI y el calendario de feriados deben provenir de una fuente autorizada (servicio del ente emisor o base interna sincronizada), inyectada al servicio, no embebida en el motor.")

# ===== 17 =====
H("17. Persistencia y auditabilidad",1)
G("Persistir, ademas de los parametros, los pagos y el ultimo resultado calculado, de modo que un credito pueda reconstruirse y auditarse. Conservar la version de los parametros regulatorios y de la serie IDI usados en cada calculo.")
G("Registrar el origen de cada IDI (publicado vs estimado) y la fecha de corte, para que el resultado sea reproducible en el tiempo.")

# ===== 18 =====
H("18. Indicadores (TIR)",1)
P("El modulo debe poder calcular la tasa interna de retorno del flujo del credito (desembolso neto y pagos), resolviendo la raiz del valor presente neto:")
formula(["VPN(r) = SUM_t  Flujo_t / (1+r)^t  = 0",
         "TIR_anual = (1 + TIR_periodo)^(periodos_por_ano) - 1"])

# ===== 19 =====
H("19. Casos limite e invariantes a garantizar",1)
table(["Caso","Comportamiento esperado"],
    [["Plazo cero","Cronograma vacio; sin cuota."],
     ["Tasa cero","Amortizacion lineal (capital / plazo)."],
     ["Sin serie IDI","Usar valor de respaldo declarado."],
     ["IDI a la baja en cancelacion","Aplicar el piso (IDI de desembolso)."],
     ["Pago parcial / en exceso","Imputacion en orden; reconduccion si corresponde."],
     ["Vencimiento en dia no habil","Mover al siguiente dia habil."],
     ["Ultima cuota","Cerrar el saldo exactamente en cero."]],
    widths=[2.3,4.2])
P("Invariantes que el desarrollo debe asegurar:")
P("Los asientos siempre cuadran (debe = haber).", bullet=True)
P("Sin prepagos, el saldo en UVC converge a cero en la ultima cuota.", bullet=True)
P("Con IDI creciente, el piso de cancelacion no altera los resultados.", bullet=True)
P("Mismas entradas producen siempre la misma salida (determinismo).", bullet=True)

# ===== 20 =====
H("20. Buenas practicas de desarrollo y pruebas",1)
P("Aislar el motor de calculo para poder probarlo sin interfaz ni base de datos.", bullet=True)
P("Validar contra tablas de referencia oficiales (reproducir un caso conocido cifra por cifra).", bullet=True)
P("Cubrir con pruebas los casos limite de la seccion 19 y los topes de cumplimiento.", bullet=True)
P("Mantener los parametros regulatorios y la serie IDI versionados y auditables.", bullet=True)
P("Documentar cada formula junto a su resultado para facilitar la revision del ente supervisor.", bullet=True)

doc.add_paragraph()
f=doc.add_paragraph(); fr=f.add_run("Documento de lineamientos y bases de diseno. Define el comportamiento esperado del modulo de credito UVC con independencia de la tecnologia o implementacion empleada.")
fr.italic=True; fr.font.size=Pt(9); fr.font.color.rgb=GRIS

doc.save(OUT)
print("OK ->", OUT)
